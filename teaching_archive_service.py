from __future__ import annotations

import hashlib
import html
import io
import json
import re
import shutil
import subprocess
import uuid
import zipfile
from datetime import date, datetime
from pathlib import Path, PurePosixPath
from typing import Any, BinaryIO

from openpyxl import load_workbook

from campus_service import MAX_UPLOAD_BYTES, CampusService, NotFound, PermissionDenied, ValidationError
from database import LearningDatabase


ARCHIVE_RECORD_TYPES = {
    "syllabus_profile", "lesson_session", "teaching_schedule_entry",
    "assessment_scheme", "assessment_component", "experiment_project",
    "experiment_material", "experiment_report", "teaching_reflection", "other",
}
IMPORT_DOCUMENT_TYPES = {
    "syllabus", "lesson_plan", "teaching_schedule", "assessment",
    "experiment_material", "experiment_report", "slides", "knowledge_graph",
    "question_bank", "other",
}
ARCHIVE_LIFECYCLES = {"processing", "review_required", "published", "withdrawn", "failed"}
ROUTING_TARGETS = {"teaching_archive", "knowledge_center", "question_center", "attachment", "ignored"}
MODERN_EXTENSIONS = {".pdf", ".docx", ".pptx", ".xlsx", ".txt", ".md"}
LEGACY_EXTENSIONS = {".doc", ".xls", ".ppt"}
ARCHIVE_EXTENSIONS = {".zip", ".rar"}
BLOCKED_EXTENSIONS = {".exe", ".dll", ".pdb", ".msi", ".com", ".scr", ".bat", ".cmd", ".ps1"}
IGNORED_EXTENSIONS = {".suo", ".cache", ".lock", ".ide", ".ide-shm", ".ide-wal", ".resources"}
IGNORED_SEGMENTS = {".vs", "bin", "obj", "backup", "__pycache__"}


def _json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _loads(value: str | None, default: Any) -> Any:
    try:
        return json.loads(value or "")
    except (TypeError, json.JSONDecodeError):
        return default


def _clean(value: Any, limit: int = 240) -> str:
    return " ".join(str(value or "").replace("\x00", "").split())[:limit]


def _safe_relative_path(value: str, fallback: str) -> str:
    normalized = str(value or fallback).replace("\\", "/").strip(" /")
    path = PurePosixPath(normalized)
    if path.is_absolute() or any(part in {"", ".", ".."} for part in path.parts):
        raise ValidationError("文件相对路径不安全")
    return "/".join(_clean(part, 120) for part in path.parts)


def _as_number(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


class TeachingArchiveService:
    """Teacher-owned teaching implementation archive, separate from the knowledge tree."""

    def __init__(self, db: LearningDatabase, campus: CampusService, ingestion: Any | None = None):
        self.db = db
        self.campus = campus
        self.ingestion = ingestion
        self.storage_root = (campus.storage_dir / "teaching_archive").resolve()
        self.storage_root.mkdir(parents=True, exist_ok=True)

    def _require_course(self, actor: dict[str, Any], course_id: str) -> dict[str, Any]:
        if actor.get("role") != "teacher" or actor.get("status") != "active":
            raise PermissionDenied("仅在职教师可以管理教学档案")
        course = self.campus.require_access(course_id, str(actor["user_id"]), "teacher")
        if course.get("course_type") != "shared_course" or course.get("owner_id") != actor["user_id"]:
            raise PermissionDenied("只能管理自己共享课程的教学档案")
        return course

    def _require_batch(self, actor: dict[str, Any], batch_id: str) -> dict[str, Any]:
        row = self.db.fetch_one(
            """SELECT b.* FROM teaching_archive_import_batches b
               JOIN courses c ON c.course_id=b.course_id
               WHERE b.batch_id=? AND b.created_by=? AND c.owner_id=?""",
            (batch_id, actor["user_id"], actor["user_id"]),
        )
        if not row:
            raise PermissionDenied("无权访问该导入批次")
        return row

    def _require_item(self, actor: dict[str, Any], item_id: str) -> dict[str, Any]:
        row = self.db.fetch_one(
            """SELECT i.*,v.course_id FROM teaching_archive_items i
               JOIN teaching_archive_versions v ON v.version_id=i.version_id
               JOIN courses c ON c.course_id=v.course_id
               WHERE i.item_id=? AND c.owner_id=?""",
            (item_id, actor["user_id"]),
        )
        if not row:
            raise PermissionDenied("无权访问该教学档案项")
        return row

    def create_import_batch(
        self, actor: dict[str, Any], course_id: str, *, term_id: str | None = None,
        batch_name: str = "", defaults: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        self._require_course(actor, course_id)
        if term_id:
            term = self.db.fetch_one(
                "SELECT term_id FROM terms WHERE term_id=? AND owner_id=?",
                (term_id, actor["user_id"]),
            )
            if not term:
                raise PermissionDenied("所选学年学期不属于当前教师")
        defaults = self._normalized_scope(defaults or {})
        batch_id = f"taib_{uuid.uuid4().hex}"
        self.db.execute(
            """INSERT INTO teaching_archive_import_batches(
                   batch_id,course_id,term_id,batch_name,defaults_json,created_by
               ) VALUES(?,?,?,?,?,?)""",
            (batch_id, course_id, term_id, _clean(batch_name, 160), _json(defaults), actor["user_id"]),
        )
        return self.get_import_batch(actor, batch_id)

    @staticmethod
    def _normalized_scope(scope: dict[str, Any]) -> dict[str, str]:
        return {
            "campus": _clean(scope.get("campus"), 100),
            "cohort_year": _clean(scope.get("cohort_year"), 32),
            "major": _clean(scope.get("major"), 120),
            "class_variant": _clean(scope.get("class_variant"), 100),
            "teaching_level": _clean(scope.get("teaching_level"), 100),
        }

    @staticmethod
    def _ignored_reason(relative_path: str, suffix: str) -> str:
        parts = {part.lower() for part in PurePosixPath(relative_path).parts}
        name = PurePosixPath(relative_path).name
        if name.startswith("~$"):
            return "office_temporary_file"
        if parts & IGNORED_SEGMENTS:
            return "build_or_backup_directory"
        if suffix in IGNORED_EXTENSIONS:
            return "runtime_or_ide_cache"
        if suffix in {".cs", ".sln", ".csproj", ".resx", ".config", ".user", ".settings"}:
            return "source_file_must_be_archived"
        return ""

    @staticmethod
    def _infer_document_type(relative_path: str, suffix: str) -> tuple[str, str]:
        text = relative_path.lower()
        if "题库" in text or "诊断库" in text or "易错" in text:
            return "question_bank", "question_center"
        if "知识图谱" in text or "知识点清单" in text:
            return "knowledge_graph", "knowledge_center"
        if "ppt" in text or suffix in {".ppt", ".pptx"}:
            if "实验" not in text:
                return "slides", "knowledge_center"
        if "教学进度" in text or "进度表" in text:
            return "teaching_schedule", "teaching_archive"
        if "教案" in text:
            return "lesson_plan", "teaching_archive"
        if "大纲" in text:
            return "syllabus", "teaching_archive"
        if "考核" in text or "成绩标准" in text:
            return "assessment", "teaching_archive"
        if "实验教学报告" in text or ("实验报告" in text and "指导" not in text):
            return "experiment_report", "teaching_archive"
        if "实验" in text:
            if suffix in ARCHIVE_EXTENSIONS or suffix in {".bak", ""}:
                return "experiment_material", "attachment"
            return "experiment_material", "teaching_archive"
        if suffix in ARCHIVE_EXTENSIONS or suffix == ".bak":
            return "experiment_material", "attachment"
        return "other", "teaching_archive"

    @staticmethod
    def _infer_scope(relative_path: str, sample_text: str = "") -> dict[str, str]:
        text = f"{relative_path} {sample_text[:5000]}"
        campus = "仁济" if "仁济" in text else ("本部" if "本部" in text else "")
        if "信息管理与信息系统" in text or "信管" in text:
            major = "信息管理与信息系统"
        elif "生物医学工程" in text or "生工" in text:
            major = "生物医学工程"
        else:
            major = ""
        cohort = ""
        match = re.search(r"(?:仁济?|本部)?\s*(20\d{2}|\d{2})\s*级?", text)
        if match:
            value = match.group(1)
            cohort = value if len(value) == 4 else f"20{value}"
        return {"campus": campus, "cohort_year": cohort, "major": major,
                "class_variant": "", "teaching_level": ""}

    def _write_stream(self, stream: BinaryIO, destination: Path) -> tuple[int, str, bytes]:
        digest = hashlib.sha256()
        size = 0
        head = bytearray()
        destination.parent.mkdir(parents=True, exist_ok=True)
        try:
            with destination.open("xb") as handle:
                while True:
                    chunk = stream.read(1024 * 1024)
                    if not chunk:
                        break
                    size += len(chunk)
                    if size > MAX_UPLOAD_BYTES:
                        raise ValidationError("单个文件超过上传大小限制")
                    if len(head) < 64:
                        head.extend(chunk[:64 - len(head)])
                    digest.update(chunk)
                    handle.write(chunk)
        except Exception:
            destination.unlink(missing_ok=True)
            raise
        if not size:
            destination.unlink(missing_ok=True)
            raise ValidationError("不能上传空文件")
        return size, digest.hexdigest(), bytes(head)

    def add_import_file(
        self, actor: dict[str, Any], batch_id: str, name: str, mime_type: str,
        stream: BinaryIO, *, relative_path: str = "",
    ) -> dict[str, Any]:
        batch = self._require_batch(actor, batch_id)
        if batch["status"] != "staging":
            raise ValidationError("该批次已提交，不能继续添加文件")
        original_name = Path(name or "document").name.replace("\x00", "")[:200]
        safe_relative = _safe_relative_path(relative_path or original_name, original_name)
        suffix = Path(original_name).suffix.lower()
        ignored_reason = self._ignored_reason(safe_relative, suffix)
        file_id = f"taif_{uuid.uuid4().hex}"
        folder = self.storage_root / batch_id
        stored = folder / f"{file_id}{suffix or '.bin'}"
        size, sha256, head = self._write_stream(stream, stored)
        risks: list[str] = []
        status = "staged"
        error = ""
        duplicate_of_document_id: str | None = None
        if ignored_reason:
            status, risks, error = "ignored", [ignored_reason], "该文件属于临时、构建或裸源码文件，已自动忽略"
        elif suffix in BLOCKED_EXTENSIONS or head[:2] == b"MZ":
            status, risks, error = "blocked", ["executable_content_blocked"], "可执行内容已拒绝，请上传去除 bin/obj 后的源码压缩包"
        elif suffix not in MODERN_EXTENSIONS | LEGACY_EXTENSIONS | ARCHIVE_EXTENSIONS | {".bak", ""}:
            status, risks, error = "blocked", ["unsupported_file_type"], "该文件类型不在教学档案安全白名单中"
        is_database_backup = head.startswith(b"TAPE")
        if suffix == "" and not is_database_backup and status == "staged":
            status, risks, error = "blocked", ["unknown_binary_type"], "无扩展名文件无法确认安全类型"
        record_type, routing_target = self._infer_document_type(safe_relative, suffix)
        if is_database_backup:
            record_type, routing_target = "experiment_material", "attachment"
        if suffix in ARCHIVE_EXTENSIONS:
            routing_target = "attachment"
        existing = self.db.fetch_one(
            """SELECT d.archive_document_id,d.original_name
                 FROM teaching_archive_documents d
                 JOIN teaching_archive_versions v USING(version_id)
                WHERE v.course_id=? AND d.sha256=? AND d.lifecycle!='withdrawn'
                ORDER BY d.created_at DESC LIMIT 1""",
            (batch["course_id"], sha256),
        )
        same_batch = self.db.fetch_one(
            """SELECT file_id,original_name FROM teaching_archive_import_files
                WHERE batch_id=? AND sha256=? ORDER BY created_at LIMIT 1""",
            (batch_id, sha256),
        )
        duplicate_action = "none"
        if existing:
            duplicate_of_document_id = str(existing["archive_document_id"])
            duplicate_action = "skip"
            risks.append("duplicate_existing_document")
            error = f"与档案库中的“{existing['original_name']}”内容相同，默认跳过；可明确选择作为新版本导入"
        elif same_batch:
            duplicate_of_document_id = f"file:{same_batch['file_id']}"
            duplicate_action = "skip"
            risks.append("duplicate_in_batch")
            error = f"与本批次的“{same_batch['original_name']}”内容相同，默认跳过"
        sample = ""
        if suffix in {".txt", ".md"} and status == "staged":
            try:
                sample = stored.read_text(encoding="utf-8-sig")[:5000]
            except UnicodeDecodeError:
                risks.append("text_encoding_requires_review")
        scope = self._normalized_scope(_loads(batch.get("defaults_json"), {}))
        inferred = self._infer_scope(safe_relative, sample)
        scope.update({key: value for key, value in inferred.items() if value})
        self.db.execute(
            """INSERT INTO teaching_archive_import_files(
                   file_id,batch_id,original_name,relative_path,stored_path,mime_type,size_bytes,
                   sha256,suggested_record_type,confirmed_record_type,routing_target,scope_json,
                   status,risk_codes_json,error_message,duplicate_of_document_id,duplicate_action
               ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (file_id, batch_id, original_name, safe_relative, str(stored), mime_type, size,
             sha256, record_type, record_type, routing_target, _json(scope), status, _json(risks), error,
             duplicate_of_document_id, duplicate_action),
        )
        self.db.execute(
            """UPDATE teaching_archive_import_batches SET file_count=(
                   SELECT COUNT(*) FROM teaching_archive_import_files WHERE batch_id=?
               ),updated_at=CURRENT_TIMESTAMP WHERE batch_id=?""", (batch_id, batch_id),
        )
        return self._file_response(self.db.fetch_one(
            "SELECT * FROM teaching_archive_import_files WHERE file_id=?", (file_id,),
        ) or {})

    @staticmethod
    def _file_response(row: dict[str, Any]) -> dict[str, Any]:
        item = dict(row)
        item.pop("stored_path", None)
        item["scope"] = _loads(item.pop("scope_json", "{}"), {})
        item["class_ids"] = _loads(item.pop("class_ids_json", "[]"), [])
        item["risk_codes"] = _loads(item.pop("risk_codes_json", "[]"), [])
        return item

    def get_import_batch(self, actor: dict[str, Any], batch_id: str) -> dict[str, Any]:
        batch = self._require_batch(actor, batch_id)
        files = self.db.fetch_all(
            "SELECT * FROM teaching_archive_import_files WHERE batch_id=? ORDER BY created_at,file_id",
            (batch_id,),
        )
        result = dict(batch)
        result["defaults"] = _loads(result.pop("defaults_json", "{}"), {})
        result["files"] = [self._file_response(row) for row in files]
        return result

    def update_import_file(
        self, actor: dict[str, Any], batch_id: str, file_id: str, payload: dict[str, Any],
    ) -> dict[str, Any]:
        batch = self._require_batch(actor, batch_id)
        if batch["status"] != "staging":
            raise ValidationError("该批次已提交")
        row = self.db.fetch_one(
            "SELECT * FROM teaching_archive_import_files WHERE file_id=? AND batch_id=?",
            (file_id, batch_id),
        )
        if not row:
            raise NotFound("导入文件不存在")
        record_type = str(payload.get("record_type", row["confirmed_record_type"]))
        route = str(payload.get("routing_target", row["routing_target"]))
        if record_type not in IMPORT_DOCUMENT_TYPES or route not in ROUTING_TARGETS:
            raise ValidationError("档案类型或路由目标无效")
        scope = self._normalized_scope(payload.get("scope") or _loads(row["scope_json"], {}))
        class_ids = list(dict.fromkeys(str(value) for value in payload.get(
            "class_ids", _loads(row["class_ids_json"], [])) if str(value)
        ))
        if class_ids:
            placeholders = ",".join("?" for _ in class_ids)
            owned = self.db.fetch_all(
                f"""SELECT class_id FROM classes WHERE class_id IN ({placeholders})
                    AND course_id=? AND teacher_id=?""",
                (*class_ids, batch["course_id"], actor["user_id"]),
            )
            if {value["class_id"] for value in owned} != set(class_ids):
                raise PermissionDenied("只能选择当前课程中自己管理的教学班")
        status = row["status"]
        duplicate_action = str(payload.get("duplicate_action", row.get("duplicate_action") or "none"))
        if duplicate_action not in {"none", "skip", "new_version"}:
            raise ValidationError("重复文件只能选择跳过或作为新版本导入")
        if status not in {"blocked", "ignored"} and payload.get("include") is False:
            status = "ignored"
        elif status == "ignored" and payload.get("include") is True:
            status = "staged"
        self.db.execute(
            """UPDATE teaching_archive_import_files SET confirmed_record_type=?,routing_target=?,
                   scope_json=?,class_ids_json=?,status=?,duplicate_action=?,updated_at=CURRENT_TIMESTAMP WHERE file_id=?""",
            (record_type, route, _json(scope), _json(class_ids), status, duplicate_action, file_id),
        )
        return self._file_response(self.db.fetch_one(
            "SELECT * FROM teaching_archive_import_files WHERE file_id=?", (file_id,),
        ) or {})

    def _find_parent_version(self, course_id: str, term_id: str | None, scope: dict[str, str]) -> str | None:
        if not scope.get("major"):
            return None
        row = self.db.fetch_one(
            """SELECT version_id FROM teaching_archive_versions
               WHERE course_id=? AND COALESCE(term_id,'')=COALESCE(?, '') AND lifecycle!='withdrawn'
                 AND scope_type IN ('course','major')
                 AND (major='' OR major=?) AND (campus='' OR campus=?)
               ORDER BY CASE scope_type WHEN 'major' THEN 0 ELSE 1 END,updated_at DESC LIMIT 1""",
            (course_id, term_id, scope.get("major", ""), scope.get("campus", "")),
        )
        return str(row["version_id"]) if row else None

    def _version_for_file(self, actor: dict[str, Any], batch: dict[str, Any], row: dict[str, Any]) -> str:
        scope = self._normalized_scope(_loads(row["scope_json"], {}))
        class_ids = _loads(row["class_ids_json"], [])
        scope_type = "class" if class_ids else ("major" if scope.get("major") else "course")
        signature = _json([batch["course_id"], batch.get("term_id"), scope_type, scope, sorted(class_ids)])
        is_new_version = row.get("duplicate_action") == "new_version"
        version_number = 1
        if is_new_version:
            latest = self.db.fetch_one(
                """SELECT MAX(version_number) latest FROM teaching_archive_versions
                    WHERE course_id=? AND COALESCE(term_id,'')=COALESCE(?, '')
                      AND scope_type=? AND campus=? AND cohort_year=? AND major=?
                      AND class_variant=? AND teaching_level=?""",
                (batch["course_id"], batch.get("term_id"), scope_type, scope["campus"],
                 scope["cohort_year"], scope["major"], scope["class_variant"], scope["teaching_level"]),
            )
            version_number = int((latest or {}).get("latest") or 0) + 1
            signature = _json([signature, row["file_id"], version_number])
        stable = hashlib.sha256(signature.encode("utf-8")).hexdigest()[:24]
        version_id = f"tav_{stable}"
        existing = self.db.fetch_one(
            "SELECT version_id FROM teaching_archive_versions WHERE version_id=?", (version_id,),
        )
        if not existing:
            labels = [scope.get("campus"), scope.get("cohort_year"), scope.get("major"),
                      scope.get("class_variant"), scope.get("teaching_level")]
            version_name = " · ".join(value for value in labels if value) or "课程通用档案"
            if is_new_version:
                version_name = f"{version_name} · v{version_number}"
            parent = self._find_parent_version(batch["course_id"], batch.get("term_id"), scope) if class_ids else None
            with self.db.connect() as conn:
                conn.execute(
                    """INSERT INTO teaching_archive_versions(
                           version_id,course_id,term_id,version_name,scope_type,campus,cohort_year,
                           major,class_variant,teaching_level,parent_version_id,created_by,version_number
                       ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (version_id, batch["course_id"], batch.get("term_id"), version_name, scope_type,
                     scope["campus"], scope["cohort_year"], scope["major"], scope["class_variant"],
                     scope["teaching_level"], parent, actor["user_id"], version_number),
                )
                conn.executemany(
                    "INSERT OR IGNORE INTO teaching_archive_version_classes(version_id,class_id) VALUES(?,?)",
                    [(version_id, class_id) for class_id in class_ids],
                )
        return version_id

    @staticmethod
    def _legacy_target_extension(suffix: str) -> str:
        return {".doc": ".docx", ".xls": ".xlsx", ".ppt": ".pptx"}[suffix]

    def _convert_legacy(self, source: Path) -> tuple[Path | None, str]:
        converter = shutil.which("soffice") or shutil.which("libreoffice")
        if not converter:
            return None, "服务器 Office 转换 Worker 不可用"
        target_ext = self._legacy_target_extension(source.suffix.lower())
        output_dir = source.parent / f"{source.stem}_converted"
        output_dir.mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [converter, "--headless", "--convert-to", target_ext.lstrip("."),
             "--outdir", str(output_dir), str(source)],
            capture_output=True, text=True, timeout=120, check=False,
        )
        output = output_dir / f"{source.stem}{target_ext}"
        if result.returncode or not output.is_file():
            return None, _clean(result.stderr or result.stdout or "旧版 Office 转换失败", 500)
        return output.resolve(), ""

    @staticmethod
    def _docx_data(source: Path) -> tuple[list[str], list[list[list[str]]]]:
        from docx import Document
        document = Document(str(source))
        paragraphs = [_clean(paragraph.text, 20000) for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[list[list[str]]] = []
        for table in document.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([_clean(cell.text, 20000) for cell in row.cells])
            tables.append(rows)
        return paragraphs, tables

    @staticmethod
    def _markdown_table(rows: list[list[str]], max_rows: int = 120) -> str:
        clean_rows = [[cell.replace("|", "\\|") for cell in row] for row in rows[:max_rows] if any(row)]
        if not clean_rows:
            return ""
        width = max(len(row) for row in clean_rows)
        clean_rows = [row + [""] * (width - len(row)) for row in clean_rows]
        return "\n".join([
            "| " + " | ".join(clean_rows[0]) + " |",
            "| " + " | ".join("---" for _ in range(width)) + " |",
            *("| " + " | ".join(row) + " |" for row in clean_rows[1:]),
        ])

    @staticmethod
    def _field_from_tables(tables: list[list[list[str]]], label: str) -> str:
        for table in tables:
            for row in table:
                for index, value in enumerate(row):
                    if label in value:
                        rest: list[str] = []
                        for cell in row[index + 1:]:
                            if cell and cell not in rest:
                                rest.append(cell)
                        if rest:
                            return "\n".join(rest)
                        if "：" in value:
                            return value.split("：", 1)[1].strip()
        return ""

    def _parse_docx(self, source: Path, document_type: str, original_name: str) -> list[dict[str, Any]]:
        paragraphs, tables = self._docx_data(source)
        full_text = "\n".join([*paragraphs, *("\n".join(" | ".join(row) for row in table) for table in tables)])
        if document_type == "lesson_plan":
            chapter = self._field_from_tables(tables, "授课章节题目")
            duration = self._field_from_tables(tables, "授课时长")
            match = re.search(r"第\s*(\d+)\s*次课", original_name)
            session_number = int(match.group(1)) if match else None
            payload = {
                "session_number": session_number,
                "chapter": chapter,
                "duration": duration,
                "audience": self._field_from_tables(tables, "授课对象"),
                "teaching_method": self._field_from_tables(tables, "教学方式"),
                "learner_analysis": self._field_from_tables(tables, "学情分析"),
                "objectives": self._field_from_tables(tables, "教学目标"),
                "key_difficulties": self._field_from_tables(tables, "教学重点"),
                "time_allocation": self._field_from_tables(tables, "课时分配"),
                "teaching_steps": tables[2] if len(tables) > 2 else [],
            }
            title = f"第{session_number}次课" if session_number else (chapter or Path(original_name).stem)
            items = [{"record_type": "lesson_session", "title": title,
                      "markdown": self._markdown_table(tables[1]) if len(tables) > 1 else full_text,
                      "payload": payload, "source": {"table": 2}, "sort": session_number or 0}]
            reflection = self._field_from_tables(tables, "教学反思")
            if reflection:
                items.append({"record_type": "teaching_reflection", "title": f"{title}教学反思",
                              "markdown": reflection, "payload": {"session_number": session_number},
                              "source": {"label": "教学反思"}, "sort": session_number or 0})
            return items
        if document_type == "syllabus":
            profile: dict[str, Any] = {}
            for label, key in (("课程名称", "course_name"), ("课程代码", "course_code"),
                               ("适用专业", "major"), ("课程学分", "credits"),
                               ("课程类型", "course_type"), ("授课对象", "audience"),
                               ("课程学时", "total_hours"), ("考核方式", "assessment_method")):
                value = self._field_from_tables(tables, label)
                if value:
                    profile[key] = value
            items = [{"record_type": "syllabus_profile", "title": profile.get("course_name") or Path(original_name).stem,
                      "markdown": "\n\n".join([*paragraphs[:80], *(self._markdown_table(t) for t in tables[:3])]),
                      "payload": profile, "source": {"tables": list(range(1, min(3, len(tables)) + 1))}, "sort": 0}]
            for table_index, table in enumerate(tables, 1):
                if not table or not any("实验项目名称" in cell for cell in table[0]):
                    continue
                for row_index, row in enumerate(table[1:], 2):
                    number = _clean(row[0]) if row else ""
                    if not re.fullmatch(r"\d+", number):
                        continue
                    title = _clean(row[1]) if len(row) > 1 else f"实验{number}"
                    items.append({"record_type": "experiment_project", "title": title,
                                  "markdown": "\n".join(value for value in row[2:] if value),
                                  "payload": {"experiment_number": int(number),
                                              "hours": row[2] if len(row) > 2 else "",
                                              "requirements": row[3] if len(row) > 3 else "",
                                              "location": row[4] if len(row) > 4 else ""},
                                  "source": {"table": table_index, "row": row_index}, "sort": int(number)})
            percentages = [float(value) for value in re.findall(r"(\d+(?:\.\d+)?)\s*%", full_text)]
            if percentages:
                unique: list[float] = []
                for value in percentages:
                    if value not in unique:
                        unique.append(value)
                likely = unique[:3]
                items.append({"record_type": "assessment_scheme", "title": "课程考核方案",
                              "markdown": self._field_from_tables(tables, "课程成绩组成"),
                              "payload": {"weights": likely, "weight_total": sum(likely)},
                              "source": {"label": "课程成绩组成"}, "sort": 900})
            return items
        if document_type in {"experiment_material", "experiment_report"}:
            match = re.search(r"实验\s*([一二三四五六七八九十\d]+)", full_text[:500] or original_name)
            number = match.group(1) if match else ""
            record = "experiment_report" if document_type == "experiment_report" else "experiment_material"
            title = paragraphs[0] if paragraphs else Path(original_name).stem
            return [{"record_type": record, "title": title,
                     "markdown": "\n\n".join([*paragraphs, *(self._markdown_table(t) for t in tables)]),
                     "payload": {"experiment_number": number,
                                 "objectives": next((p for p in paragraphs if "实验目" in p), "")},
                     "source": {"document": original_name}, "sort": int(number) if number.isdigit() else 0}]
        return [{"record_type": "other", "title": Path(original_name).stem,
                 "markdown": "\n\n".join([*paragraphs, *(self._markdown_table(t) for t in tables)]),
                 "payload": {}, "source": {"document": original_name}, "sort": 0}]

    @staticmethod
    def _cell_value(value: Any) -> str:
        if isinstance(value, datetime):
            return value.date().isoformat()
        if isinstance(value, date):
            return value.isoformat()
        return _clean(value, 2000)

    def _parse_schedule(self, source: Path) -> list[dict[str, Any]]:
        if source.suffix.lower() == ".xls":
            return self._parse_legacy_schedule(source)
        workbook = load_workbook(source, read_only=True, data_only=True)
        items: list[dict[str, Any]] = []
        order = 0
        for sheet in workbook.worksheets:
            rows = [[self._cell_value(value) for value in row] for row in sheet.iter_rows(values_only=True)]
            header_index = next((index for index, row in enumerate(rows)
                                 if "周次" in row and "课程名称" in row), None)
            if header_index is None:
                continue
            header = rows[header_index]
            normalized = [value.replace("\n", "") for value in header]
            for row_number, row in enumerate(rows[header_index + 1:], header_index + 2):
                if not row or not row[0] or not re.match(r"^\d+$", row[0]):
                    continue
                data = {normalized[index] or f"column_{index + 1}": row[index]
                        for index in range(min(len(normalized), len(row))) if row[index]}
                order += 1
                title = data.get("理论/实验授课内容") or data.get("授课内容") or f"第{row[0]}周"
                items.append({"record_type": "teaching_schedule_entry", "title": title,
                              "markdown": " · ".join(f"{key}：{value}" for key, value in data.items()),
                              "payload": {"sheet": sheet.title, "row": row_number, **data},
                              "source": {"sheet": sheet.title, "row": row_number}, "sort": order})
        workbook.close()
        if not items:
            raise ValidationError("未在教学进度表中找到包含“周次”和“课程名称”的表头")
        return items

    def _parse_legacy_schedule(self, source: Path) -> list[dict[str, Any]]:
        """Read old Excel schedules deterministically, even when preview conversion is offline."""
        import xlrd

        workbook = xlrd.open_workbook(str(source), on_demand=True)
        items: list[dict[str, Any]] = []
        order = 0
        for sheet in workbook.sheets():
            rows = [[self._cell_value(sheet.cell_value(row, column))
                     for column in range(sheet.ncols)] for row in range(sheet.nrows)]
            header_index = next((index for index, row in enumerate(rows)
                                 if "周次" in row and "课程名称" in row), None)
            if header_index is None:
                continue
            header = [value.replace("\n", "") for value in rows[header_index]]
            for row_number, row in enumerate(rows[header_index + 1:], header_index + 2):
                week = row[0].removesuffix(".0") if row else ""
                if not re.fullmatch(r"\d+", week):
                    continue
                data = {header[index] or f"column_{index + 1}": row[index]
                        for index in range(min(len(header), len(row))) if row[index]}
                data["周次"] = week
                order += 1
                title = data.get("理论/实验授课内容") or data.get("授课内容") or f"第{week}周"
                items.append({"record_type": "teaching_schedule_entry", "title": title,
                              "markdown": " · ".join(f"{key}：{value}" for key, value in data.items()),
                              "payload": {"sheet": sheet.name, "row": row_number, **data},
                              "source": {"sheet": sheet.name, "row": row_number}, "sort": order})
        workbook.release_resources()
        if not items:
            raise ValidationError("未在旧版教学进度表中找到包含“周次”和“课程名称”的表头")
        return items

    @staticmethod
    def _parse_text(source: Path, document_type: str, original_name: str) -> list[dict[str, Any]]:
        try:
            text = source.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ValidationError("文本文件必须使用 UTF-8 编码") from exc
        if document_type == "assessment":
            components = []
            for label, value in re.findall(
                r"(平时成绩|阶段考核成绩|期末成绩)[^\d%]{0,20}(\d+(?:\.\d+)?)\s*%", text
            ):
                if not any(item["name"] == label for item in components):
                    components.append({"name": label, "weight": float(value)})
            total = sum(item["weight"] for item in components)
            return [
                {"record_type": "assessment_scheme", "title": "课程考核方案",
                 "markdown": text, "payload": {"components": components, "weight_total": total},
                 "source": {"document": original_name}, "sort": 0},
                *({"record_type": "assessment_component", "title": component["name"],
                   "markdown": f"{component['name']}：{component['weight']:g}%", "payload": component,
                   "source": {"document": original_name, "label": component["name"]}, "sort": index + 1}
                  for index, component in enumerate(components)),
            ]
        record = "experiment_report" if document_type == "experiment_report" else (
            "experiment_material" if document_type == "experiment_material" else "other"
        )
        number = TeachingArchiveService._experiment_number(f"{original_name}\n{text[:500]}")
        return [{"record_type": record, "title": Path(original_name).stem, "markdown": text,
                 "payload": {"experiment_number": number},
                 "source": {"document": original_name}, "sort": number or 0}]

    @staticmethod
    def _experiment_number(value: Any) -> int | None:
        """Extract a conservative experiment number without guessing from unrelated digits."""
        text = str(value or "")
        match = re.search(r"实验\s*([一二三四五六七八九十百\d]+)", text, re.IGNORECASE)
        if not match:
            return None
        token = match.group(1)
        if token.isdigit():
            return int(token)
        digits = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                  "六": 6, "七": 7, "八": 8, "九": 9}
        if token == "十":
            return 10
        if "十" in token:
            left, right = token.split("十", 1)
            return (digits.get(left, 1) * 10) + digits.get(right, 0)
        return digits.get(token)

    def _link_experiment_attachments(self, course_id: str) -> None:
        """Persist only deterministic attachment-to-project matches within one teaching version."""
        projects = self.db.fetch_all(
            """SELECT i.item_id,i.version_id,i.title,i.structured_json
                 FROM teaching_archive_items i JOIN teaching_archive_versions v USING(version_id)
                WHERE v.course_id=? AND i.record_type='experiment_project'
                  AND i.lifecycle!='withdrawn'""", (course_id,),
        )
        by_version: dict[str, dict[int, list[str]]] = {}
        for project in projects:
            payload = _loads(project["structured_json"], {})
            number = self._experiment_number(f"实验{payload.get('experiment_number', '')} {project['title']}")
            if number is not None:
                by_version.setdefault(str(project["version_id"]), {}).setdefault(number, []).append(
                    str(project["item_id"])
                )
        attachments = self.db.fetch_all(
            """SELECT a.attachment_id,a.version_id,a.original_name,f.relative_path
                 FROM teaching_archive_attachments a
                 JOIN teaching_archive_import_files f ON f.file_id=a.file_id
                 JOIN teaching_archive_versions v ON v.version_id=a.version_id
                WHERE v.course_id=?""", (course_id,),
        )
        for attachment in attachments:
            number = self._experiment_number(
                f"{attachment.get('relative_path') or ''} {attachment['original_name']}"
            )
            candidates = by_version.get(str(attachment["version_id"]), {}).get(number or -1, [])
            # Ambiguous projects deliberately remain unlinked for teacher review.
            target = candidates[0] if len(candidates) == 1 else None
            self.db.execute(
                "UPDATE teaching_archive_attachments SET experiment_item_id=? WHERE attachment_id=?",
                (target, attachment["attachment_id"]),
            )

    def _parse_document(self, source: Path, document_type: str, original_name: str) -> list[dict[str, Any]]:
        suffix = source.suffix.lower()
        if suffix == ".docx":
            return self._parse_docx(source, document_type, original_name)
        if suffix in {".xlsx", ".xls"} and document_type == "teaching_schedule":
            return self._parse_schedule(source)
        if suffix in {".txt", ".md"}:
            return self._parse_text(source, document_type, original_name)
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(source))
            pages = [page.extract_text() or "" for page in reader.pages]
            return [{"record_type": "other", "title": Path(original_name).stem,
                     "markdown": "\n\n".join(pages), "payload": {"page_count": len(pages)},
                     "source": {"pages": list(range(1, len(pages) + 1))}, "sort": 0}]
        return []

    @staticmethod
    def _archive_manifest(source: Path) -> tuple[list[dict[str, Any]], list[str]]:
        risks: list[str] = []
        entries: list[dict[str, Any]] = []
        if source.suffix.lower() == ".zip":
            with zipfile.ZipFile(source) as archive:
                infos = archive.infolist()
                if len(infos) > 2000:
                    raise ValidationError("压缩包文件数量超过 2000")
                total = 0
                for info in infos:
                    path = PurePosixPath(info.filename.replace("\\", "/"))
                    if path.is_absolute() or ".." in path.parts:
                        raise ValidationError("压缩包包含路径穿越条目")
                    total += int(info.file_size)
                    if total > 2 * 1024 * 1024 * 1024:
                        raise ValidationError("压缩包解压后总大小超过 2GB")
                    suffix = Path(path.name).suffix.lower()
                    if suffix in BLOCKED_EXTENSIONS:
                        risks.append("archive_contains_executable")
                    if suffix in ARCHIVE_EXTENSIONS:
                        risks.append("nested_archive_requires_review")
                    entries.append({"path": str(path), "size": int(info.file_size), "directory": info.is_dir()})
        else:
            tar = shutil.which("tar")
            if not tar:
                return [], ["rar_manifest_unavailable"]
            result = subprocess.run([tar, "-tf", str(source)], capture_output=True, text=True,
                                    timeout=30, check=False)
            if result.returncode:
                return [], ["rar_manifest_unavailable"]
            for value in result.stdout.splitlines()[:2001]:
                path = PurePosixPath(value.replace("\\", "/"))
                if path.is_absolute() or ".." in path.parts:
                    raise ValidationError("压缩包包含路径穿越条目")
                entries.append({"path": str(path), "size": None, "directory": value.endswith("/")})
            if len(entries) > 2000:
                raise ValidationError("压缩包文件数量超过 2000")
        return entries, sorted(set(risks))

    @staticmethod
    def _preview_kind(source: Path) -> str:
        return {".pdf": "pdf", ".docx": "docx", ".pptx": "pptx", ".xlsx": "xlsx",
                ".txt": "text", ".md": "markdown"}.get(source.suffix.lower(), "unavailable")

    def _route_external(self, actor: dict[str, Any], batch: dict[str, Any], row: dict[str, Any]) -> str | None:
        suffix = Path(row["stored_path"]).suffix.lower()
        if row["routing_target"] == "knowledge_center" and self.ingestion and suffix in MODERN_EXTENSIONS - {".xlsx"}:
            with Path(row["stored_path"]).open("rb") as stream:
                result = self.ingestion.queue_document_stream(
                    actor, batch["course_id"], row["original_name"], row["mime_type"], stream,
                    analysis_mode="local",
                )
            return str(result["document_id"])
        return None

    def _insert_items(
        self, archive_document_id: str, version_id: str, items: list[dict[str, Any]],
        document_risks: list[str],
    ) -> tuple[int, int]:
        published = review = 0
        with self.db.connect() as conn:
            for index, item in enumerate(items):
                risks = list(document_risks)
                payload = item.get("payload") or {}
                if not item.get("source"):
                    risks.append("missing_source")
                if item["record_type"] == "assessment_scheme":
                    total = _as_number(payload.get("weight_total"))
                    if total is not None and abs(total - 100) > 0.01:
                        risks.append("assessment_weight_mismatch")
                lifecycle = "review_required" if risks else "published"
                published += lifecycle == "published"
                review += lifecycle == "review_required"
                conn.execute(
                    """INSERT INTO teaching_archive_items(
                           item_id,archive_document_id,version_id,record_type,title,content_markdown,
                           structured_json,source_json,lifecycle,risk_codes_json,sort_order
                       ) VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                    (f"tai_{uuid.uuid4().hex}", archive_document_id, version_id,
                     item["record_type"], _clean(item.get("title"), 240) or "未命名档案项",
                     str(item.get("markdown") or ""), _json(payload), _json(item.get("source") or {}),
                     lifecycle, _json(sorted(set(risks))), int(item.get("sort", index))),
                )
        return published, review

    def commit_import_batch(self, actor: dict[str, Any], batch_id: str) -> dict[str, Any]:
        batch = self._require_batch(actor, batch_id)
        if batch["status"] != "staging":
            return self.get_import_batch(actor, batch_id)
        self.db.execute(
            "UPDATE teaching_archive_import_batches SET status='committing',updated_at=CURRENT_TIMESTAMP WHERE batch_id=?",
            (batch_id,),
        )
        files = self.db.fetch_all(
            "SELECT * FROM teaching_archive_import_files WHERE batch_id=? ORDER BY created_at,file_id", (batch_id,),
        )
        published_count = review_count = error_count = 0
        for row in files:
            if row["status"] in {"ignored", "blocked"}:
                error_count += row["status"] == "blocked"
                continue
            if row.get("duplicate_action") == "skip":
                self.db.execute(
                    "UPDATE teaching_archive_import_files SET status='ignored',updated_at=CURRENT_TIMESTAMP WHERE file_id=?",
                    (row["file_id"],),
                )
                continue
            try:
                version_id = self._version_for_file(actor, batch, row)
                source = Path(row["stored_path"]).resolve()
                if self.storage_root not in source.parents or not source.is_file():
                    raise ValidationError("导入文件存储位置无效")
                document_risks = list(_loads(row["risk_codes_json"], []))
                course_document_id = self._route_external(actor, batch, row)
                if row["routing_target"] in {"knowledge_center", "question_center"}:
                    if not course_document_id:
                        document_risks.append("manual_route_required")
                    self.db.execute(
                        "UPDATE teaching_archive_import_files SET status='routed',risk_codes_json=?,updated_at=CURRENT_TIMESTAMP WHERE file_id=?",
                        (_json(sorted(set(document_risks))), row["file_id"]),
                    )
                    continue
                archive_document_id = f"tad_{uuid.uuid4().hex}"
                if row["routing_target"] == "attachment":
                    manifest: list[dict[str, Any]] = []
                    if source.suffix.lower() in ARCHIVE_EXTENSIONS:
                        manifest, archive_risks = self._archive_manifest(source)
                        document_risks.extend(archive_risks)
                    with source.open("rb") as handle:
                        attachment_type = "database_backup" if handle.read(4) == b"TAPE" else "archive"
                    lifecycle = "review_required" if document_risks else "published"
                    self.db.execute(
                        """INSERT INTO teaching_archive_attachments(
                               attachment_id,file_id,version_id,attachment_type,original_name,stored_path,
                               sha256,manifest_json,lifecycle
                           ) VALUES(?,?,?,?,?,?,?,?,?)""",
                        (f"taa_{uuid.uuid4().hex}", row["file_id"], version_id, attachment_type,
                         row["original_name"], str(source), row["sha256"], _json(manifest), lifecycle),
                    )
                    published_count += lifecycle == "published"
                    review_count += lifecycle == "review_required"
                    self.db.execute(
                        "UPDATE teaching_archive_import_files SET status=?,risk_codes_json=?,updated_at=CURRENT_TIMESTAMP WHERE file_id=?",
                        (lifecycle, _json(sorted(set(document_risks))), row["file_id"]),
                    )
                    continue
                parse_source = source
                preview_source = source
                conversion_status = "ready"
                if source.suffix.lower() in LEGACY_EXTENSIONS:
                    converted, error = self._convert_legacy(source)
                    if not converted:
                        document_risks.append("legacy_conversion_unavailable")
                        conversion_status = "unavailable"
                        if source.suffix.lower() == ".xls" and row["confirmed_record_type"] == "teaching_schedule":
                            parse_source = source
                    else:
                        parse_source = converted
                        preview_source = converted
                can_parse = conversion_status == "ready" or (
                    source.suffix.lower() == ".xls" and row["confirmed_record_type"] == "teaching_schedule"
                )
                items = self._parse_document(parse_source, row["confirmed_record_type"], row["original_name"]) \
                    if can_parse else []
                if not items:
                    document_risks.append("no_structured_content")
                if any(
                    item["record_type"] == "assessment_scheme"
                    and (total := _as_number((item.get("payload") or {}).get("weight_total"))) is not None
                    and abs(total - 100) > 0.01
                    for item in items
                ):
                    document_risks.append("assessment_weight_mismatch")
                lifecycle = "review_required" if document_risks else "published"
                self.db.execute(
                    """INSERT INTO teaching_archive_documents(
                           archive_document_id,file_id,version_id,course_document_id,record_type,
                           original_name,relative_path,stored_path,sha256,preview_kind,preview_path,
                           conversion_status,lifecycle,risk_codes_json
                       ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (archive_document_id, row["file_id"], version_id, course_document_id,
                     row["confirmed_record_type"], row["original_name"], row["relative_path"],
                     str(source), row["sha256"], self._preview_kind(preview_source),
                     str(preview_source) if preview_source != source else "", conversion_status,
                     lifecycle, _json(sorted(set(document_risks)))),
                )
                published, review = self._insert_items(
                    archive_document_id, version_id, items, document_risks,
                )
                published_count += published
                review_count += review or (1 if not items else 0)
                self.db.execute(
                    "UPDATE teaching_archive_import_files SET status=?,risk_codes_json=?,updated_at=CURRENT_TIMESTAMP WHERE file_id=?",
                    (lifecycle, _json(sorted(set(document_risks))), row["file_id"]),
                )
            except Exception as exc:
                error_count += 1
                self.db.execute(
                    """UPDATE teaching_archive_import_files SET status='failed',error_message=?,
                       risk_codes_json=?,updated_at=CURRENT_TIMESTAMP WHERE file_id=?""",
                    (_clean(exc, 500), _json(["processing_failed"]), row["file_id"]),
                )
        self._link_experiment_attachments(batch["course_id"])
        self._validate_sequences(batch["course_id"])
        self._validate_hours(batch["course_id"])
        status = "completed_with_warnings" if review_count or error_count else "completed"
        self.db.execute(
            """UPDATE teaching_archive_import_batches SET status=?,published_count=?,review_count=?,
               error_count=?,updated_at=CURRENT_TIMESTAMP WHERE batch_id=?""",
            (status, published_count, review_count, error_count, batch_id),
        )
        return self.get_import_batch(actor, batch_id)

    def _validate_sequences(self, course_id: str) -> None:
        for record_type, risk_code, key in (
            ("lesson_session", "lesson_sequence_gap", "session_number"),
            ("experiment_project", "experiment_sequence_gap", "experiment_number"),
        ):
            rows = self.db.fetch_all(
                """SELECT i.item_id,i.version_id,i.structured_json,i.risk_codes_json
                   FROM teaching_archive_items i JOIN teaching_archive_versions v USING(version_id)
                   WHERE v.course_id=? AND i.record_type=? AND i.lifecycle!='withdrawn'""",
                (course_id, record_type),
            )
            grouped: dict[str, list[tuple[str, int, list[str]]]] = {}
            for row in rows:
                number = _loads(row["structured_json"], {}).get(key)
                try:
                    parsed = int(number)
                except (TypeError, ValueError):
                    continue
                grouped.setdefault(row["version_id"], []).append(
                    (row["item_id"], parsed, _loads(row["risk_codes_json"], []))
                )
            for values in grouped.values():
                numbers = [value[1] for value in values]
                expected = list(range(min(numbers), max(numbers) + 1)) if numbers else []
                if len(numbers) != len(set(numbers)) or sorted(set(numbers)) != expected:
                    for item_id, _number, risks in values:
                        if risk_code not in risks:
                            risks.append(risk_code)
                        self.db.execute(
                            """UPDATE teaching_archive_items SET lifecycle='review_required',risk_codes_json=?,
                            updated_at=CURRENT_TIMESTAMP WHERE item_id=?""", (_json(risks), item_id),
                        )

    @staticmethod
    def _first_number(value: Any) -> float | None:
        match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
        return float(match.group()) if match else None

    @staticmethod
    def _schedule_hours(payload: dict[str, Any]) -> float | None:
        for key, value in payload.items():
            if "学时" not in key:
                continue
            numbers = [float(number) for number in re.findall(r"\d+(?:\.\d+)?", str(value))]
            if numbers:
                return sum(numbers)
        return None

    def _validate_hours(self, course_id: str) -> None:
        rows = self.db.fetch_all(
            """SELECT i.item_id,i.version_id,i.record_type,i.structured_json,i.risk_codes_json
                 FROM teaching_archive_items i JOIN teaching_archive_versions v USING(version_id)
                WHERE v.course_id=? AND i.record_type IN(
                      'syllabus_profile','lesson_session','teaching_schedule_entry')""", (course_id,),
        )
        grouped: dict[str, list[dict[str, Any]]] = {}
        for row in rows:
            grouped.setdefault(str(row["version_id"]), []).append(row)
        for values in grouped.values():
            syllabus = next((self._first_number(_loads(row["structured_json"], {}).get("total_hours"))
                             for row in values if row["record_type"] == "syllabus_profile"), None)
            if syllabus is None:
                continue
            schedule_hours = [self._schedule_hours(_loads(row["structured_json"], {}))
                              for row in values if row["record_type"] == "teaching_schedule_entry"]
            lesson_hours = [self._first_number(_loads(row["structured_json"], {}).get("duration"))
                            for row in values if row["record_type"] == "lesson_session"]
            comparable = [sum(value for value in totals if value is not None)
                          for totals in (schedule_hours, lesson_hours) if any(value is not None for value in totals)]
            if not comparable or all(abs(total - syllabus) <= 0.01 for total in comparable):
                continue
            for row in values:
                risks = _loads(row["risk_codes_json"], [])
                if "hours_mismatch" not in risks:
                    risks.append("hours_mismatch")
                self.db.execute(
                    """UPDATE teaching_archive_items SET lifecycle='review_required',risk_codes_json=?,
                       updated_at=CURRENT_TIMESTAMP WHERE item_id=?""", (_json(risks), row["item_id"]),
                )

    @staticmethod
    def _decode_item(row: dict[str, Any]) -> dict[str, Any]:
        item = dict(row)
        item["structured"] = _loads(item.pop("structured_json", "{}"), {})
        item["source"] = _loads(item.pop("source_json", "{}"), {})
        item["risk_codes"] = _loads(item.pop("risk_codes_json", "[]"), [])
        return item

    def _backfill_legacy_archive(self, actor: dict[str, Any], course_id: str) -> None:
        """Lazily copy already-approved syllabus archive nodes into the dedicated model."""
        legacy = self.db.fetch_all(
            """SELECT n.*,d.original_name,d.stored_path,d.sha256 AS document_sha256
               FROM knowledge_nodes n JOIN course_documents d ON d.document_id=n.document_id
               WHERE n.course_id=? AND n.content_domain='teaching_archive'
                 AND n.node_type='knowledge_point' AND n.status!='rejected'
                 AND NOT EXISTS(
                     SELECT 1 FROM teaching_archive_documents ad WHERE ad.course_document_id=n.document_id
                 ) ORDER BY n.document_id,n.sort_order""", (course_id,),
        )
        if not legacy:
            return
        version = self.db.fetch_one(
            """SELECT version_id FROM teaching_archive_versions WHERE course_id=?
               ORDER BY CASE WHEN version_id LIKE 'tav_legacy_%' THEN 0 ELSE 1 END,created_at LIMIT 1""",
            (course_id,),
        )
        if not version:
            version_id = f"tav_legacy_{hashlib.sha256(course_id.encode()).hexdigest()[:20]}"
            self.db.execute(
                """INSERT INTO teaching_archive_versions(
                       version_id,course_id,version_name,scope_type,created_by,lifecycle
                   ) VALUES(?,?,'历史大纲档案','course',?,'published')""",
                (version_id, course_id, actor["user_id"]),
            )
        else:
            version_id = str(version["version_id"])
        by_document: dict[str, list[dict[str, Any]]] = {}
        for row in legacy:
            by_document.setdefault(str(row["document_id"]), []).append(row)
        batch_id = f"taib_legacy_{hashlib.sha256(course_id.encode()).hexdigest()[:20]}"
        if not self.db.fetch_one("SELECT batch_id FROM teaching_archive_import_batches WHERE batch_id=?", (batch_id,)):
            self.db.execute(
                """INSERT INTO teaching_archive_import_batches(
                       batch_id,course_id,batch_name,status,file_count,published_count,created_by
                   ) VALUES(?,?,'历史大纲迁移','completed',?,0,?)""",
                (batch_id, course_id, len(by_document), actor["user_id"]),
            )
        category_types = {
            "course_profile": "syllabus_profile", "objectives": "syllabus_profile",
            "teaching_design": "lesson_session", "assessment": "assessment_scheme",
        }
        for document_id, rows in by_document.items():
            first = rows[0]
            file_id = f"taif_legacy_{hashlib.sha256(document_id.encode()).hexdigest()[:20]}"
            archive_document_id = f"tad_legacy_{hashlib.sha256(document_id.encode()).hexdigest()[:20]}"
            source = Path(str(first["stored_path"]))
            sha256 = str(first.get("document_sha256") or hashlib.sha256(document_id.encode()).hexdigest())
            with self.db.connect() as conn:
                conn.execute(
                    """INSERT OR IGNORE INTO teaching_archive_import_files(
                           file_id,batch_id,original_name,relative_path,stored_path,size_bytes,sha256,
                           suggested_record_type,confirmed_record_type,routing_target,status
                       ) VALUES(?,?,?,?,?,?,?,'syllabus','syllabus','teaching_archive','published')""",
                    (file_id, batch_id, first["original_name"], first["original_name"],
                     str(source), source.stat().st_size if source.is_file() else 0, sha256),
                )
                conn.execute(
                    """INSERT OR IGNORE INTO teaching_archive_documents(
                           archive_document_id,file_id,version_id,course_document_id,record_type,
                           original_name,relative_path,stored_path,sha256,preview_kind,lifecycle
                       ) VALUES(?,?,?,?,?,?,?,?,?,?, 'published')""",
                    (archive_document_id, file_id, version_id, document_id, "syllabus",
                     first["original_name"], first["original_name"], str(source), sha256,
                     self._preview_kind(source)),
                )
                for index, row in enumerate(rows):
                    record_type = category_types.get(str(row.get("teaching_category") or ""), "other")
                    conn.execute(
                        """INSERT OR IGNORE INTO teaching_archive_items(
                               item_id,archive_document_id,version_id,record_type,title,content_markdown,
                               structured_json,source_json,lifecycle,sort_order
                           ) VALUES(?,?,?,?,?,?, '{}',?,'published',?)""",
                        (f"tai_legacy_{hashlib.sha256(str(row['node_id']).encode()).hexdigest()[:20]}",
                         archive_document_id, version_id, record_type, row["title"], row["markdown"],
                         _json({"knowledge_node_id": row["node_id"],
                                "source_pages": _loads(row.get("source_pages_json"), [])}), index),
                    )

    def workbench(self, actor: dict[str, Any], course_id: str, filters: dict[str, str] | None = None) -> dict[str, Any]:
        course = self._require_course(actor, course_id)
        self._backfill_legacy_archive(actor, course_id)
        filters = {key: _clean(value, 120) for key, value in (filters or {}).items() if value}
        conditions = ["v.course_id=?"]
        params: list[Any] = [course_id]
        for key in ("campus", "cohort_year", "major", "teaching_level", "class_variant", "lifecycle"):
            if filters.get(key):
                conditions.append(f"v.{key}=?")
                params.append(filters[key])
        if filters.get("term_id"):
            conditions.append("v.term_id=?")
            params.append(filters["term_id"])
        if filters.get("class_id"):
            conditions.append("EXISTS(SELECT 1 FROM teaching_archive_version_classes vc WHERE vc.version_id=v.version_id AND vc.class_id=?)")
            params.append(filters["class_id"])
        versions = self.db.fetch_all(
            f"""SELECT v.*,t.term_name,t.academic_year,t.teaching_period,
                       GROUP_CONCAT(vc.class_id) class_ids
                  FROM teaching_archive_versions v
                  LEFT JOIN terms t ON t.term_id=v.term_id
                  LEFT JOIN teaching_archive_version_classes vc ON vc.version_id=v.version_id
                 WHERE {' AND '.join(conditions)} GROUP BY v.version_id ORDER BY
                       COALESCE(NULLIF(t.academic_year,''),t.term_name) DESC,v.campus,v.cohort_year,v.major,v.version_name""",
            tuple(params),
        )
        version_ids = [row["version_id"] for row in versions]
        items: list[dict[str, Any]] = []
        documents: list[dict[str, Any]] = []
        attachments: list[dict[str, Any]] = []
        if version_ids:
            placeholders = ",".join("?" for _ in version_ids)
            item_conditions = [f"i.version_id IN ({placeholders})"]
            item_params: list[Any] = list(version_ids)
            if filters.get("record_type"):
                item_conditions.append("i.record_type=?")
                item_params.append(filters["record_type"])
            if filters.get("status"):
                item_conditions.append("i.lifecycle=?")
                item_params.append(filters["status"])
            items = [self._decode_item(row) for row in self.db.fetch_all(
                f"""SELECT i.*,d.original_name,d.preview_kind,d.conversion_status
                    FROM teaching_archive_items i JOIN teaching_archive_documents d USING(archive_document_id)
                    WHERE {' AND '.join(item_conditions)} ORDER BY i.record_type,i.sort_order,i.created_at""",
                tuple(item_params),
            )]
            documents = self.db.fetch_all(
                f"""SELECT d.* FROM teaching_archive_documents d
                    WHERE d.version_id IN ({placeholders}) ORDER BY d.created_at DESC""", tuple(version_ids),
            )
            attachments = self.db.fetch_all(
                f"""SELECT a.*,f.relative_path FROM teaching_archive_attachments a
                    JOIN teaching_archive_import_files f ON f.file_id=a.file_id
                    WHERE a.version_id IN ({placeholders}) ORDER BY a.created_at DESC""", tuple(version_ids),
            )
        for row in versions:
            row["class_ids"] = [value for value in str(row.get("class_ids") or "").split(",") if value]
        for row in documents:
            row.pop("stored_path", None)
            row.pop("preview_path", None)
            row["risk_codes"] = _loads(row.pop("risk_codes_json", "[]"), [])
        for row in attachments:
            row.pop("stored_path", None)
            row["manifest"] = _loads(row.pop("manifest_json", "[]"), [])
        classes = self.db.fetch_all(
            """SELECT cl.*,t.term_name,t.academic_year,t.teaching_period,
                      (SELECT COUNT(*) FROM class_memberships m WHERE m.class_id=cl.class_id AND m.status='active') member_count
               FROM classes cl JOIN terms t ON t.term_id=cl.term_id
               WHERE cl.course_id=? AND cl.teacher_id=? ORDER BY t.academic_year DESC,t.teaching_period,
                    cl.campus,cl.cohort_year,cl.major,cl.class_name""",
            (course_id, actor["user_id"]),
        )
        published_docs = sum(row["lifecycle"] == "published" for row in documents)
        review_items = sum(item["lifecycle"] == "review_required" for item in items)
        required = {"syllabus_profile", "lesson_session", "teaching_schedule_entry", "assessment_scheme", "experiment_project"}
        present = {item["record_type"] for item in items if item["lifecycle"] == "published"}
        covered_classes = {class_id for version in versions for class_id in version["class_ids"]}
        metrics = {
            "completeness": round(100 * len(required & present) / len(required)),
            "published_documents": published_docs,
            "review_required": review_items,
            "class_coverage": round(100 * len(covered_classes) / len(classes)) if classes else 0,
            "hours_consistent": not any("hours_mismatch" in item["risk_codes"] for item in items),
        }
        return {"course": course, "metrics": metrics, "versions": versions, "classes": classes,
                "items": items, "documents": documents, "attachments": attachments,
                "record_types": sorted(ARCHIVE_RECORD_TYPES), "filters": filters}

    def update_item(self, actor: dict[str, Any], item_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        item = self._require_item(actor, item_id)
        target_version = str(payload.get("target_version_id") or item["version_id"])
        if target_version != item["version_id"]:
            version = self.db.fetch_one(
                """SELECT v.version_id FROM teaching_archive_versions v JOIN courses c USING(course_id)
                   WHERE v.version_id=? AND c.owner_id=?""", (target_version, actor["user_id"]),
            )
            if not version:
                raise PermissionDenied("无权覆盖目标教学版本")
            patch = {key: payload[key] for key in ("title", "content_markdown", "structured") if key in payload}
            override_id = f"taio_{uuid.uuid4().hex}"
            self.db.execute(
                """INSERT INTO teaching_archive_item_overrides(override_id,base_item_id,version_id,patch_json,created_by)
                   VALUES(?,?,?,?,?) ON CONFLICT(base_item_id,version_id) DO UPDATE SET
                   patch_json=excluded.patch_json,updated_at=CURRENT_TIMESTAMP""",
                (override_id, item_id, target_version, _json(patch), actor["user_id"]),
            )
            return {"item_id": item_id, "version_id": target_version, "inheritance_state": "overridden", **patch}
        title = _clean(payload.get("title", item["title"]), 240)
        markdown = str(payload.get("content_markdown", item["content_markdown"]))
        structured = payload.get("structured", _loads(item["structured_json"], {}))
        self.db.execute(
            """UPDATE teaching_archive_items SET title=?,content_markdown=?,structured_json=?,
               updated_at=CURRENT_TIMESTAMP WHERE item_id=?""", (title, markdown, _json(structured), item_id),
        )
        return self._decode_item(self._require_item(actor, item_id))

    def set_item_lifecycle(self, actor: dict[str, Any], item_id: str, lifecycle: str) -> dict[str, Any]:
        self._require_item(actor, item_id)
        if lifecycle not in {"published", "withdrawn"}:
            raise ValidationError("只能发布或撤回教学档案项")
        self.db.execute(
            "UPDATE teaching_archive_items SET lifecycle=?,updated_at=CURRENT_TIMESTAMP WHERE item_id=?",
            (lifecycle, item_id),
        )
        return self._decode_item(self._require_item(actor, item_id))

    def delete_document(self, actor: dict[str, Any], archive_document_id: str) -> dict[str, Any]:
        row = self.db.fetch_one(
            """SELECT d.archive_document_id,d.file_id,d.version_id,d.original_name,d.stored_path,d.preview_path
                 FROM teaching_archive_documents d
                 JOIN teaching_archive_versions v USING(version_id)
                 JOIN courses c USING(course_id)
                WHERE d.archive_document_id=? AND c.owner_id=?""",
            (archive_document_id, actor["user_id"]),
        )
        if not row:
            raise PermissionDenied("无权删除该教学档案文件")
        self._remove_archive_file(row)
        return {"deleted": archive_document_id, "original_name": row["original_name"]}

    def delete_attachment(self, actor: dict[str, Any], attachment_id: str) -> dict[str, Any]:
        row = self.db.fetch_one(
            """SELECT a.attachment_id,a.file_id,a.version_id,a.original_name,a.stored_path,'' preview_path
                 FROM teaching_archive_attachments a
                 JOIN teaching_archive_versions v USING(version_id)
                 JOIN courses c USING(course_id)
                WHERE a.attachment_id=? AND c.owner_id=?""",
            (attachment_id, actor["user_id"]),
        )
        if not row:
            raise PermissionDenied("无权删除该教学档案附件")
        self._remove_archive_file(row)
        return {"deleted": attachment_id, "original_name": row["original_name"]}

    def _remove_archive_file(self, row: dict[str, Any]) -> None:
        with self.db.connect() as conn:
            conn.execute("DELETE FROM teaching_archive_import_files WHERE file_id=?", (row["file_id"],))
            conn.execute(
                """DELETE FROM teaching_archive_versions WHERE version_id=?
                     AND NOT EXISTS(SELECT 1 FROM teaching_archive_documents WHERE version_id=?)
                     AND NOT EXISTS(SELECT 1 FROM teaching_archive_attachments WHERE version_id=?)""",
                (row["version_id"], row["version_id"], row["version_id"]),
            )
        for value in (row.get("preview_path"), row.get("stored_path")):
            if not value:
                continue
            path = Path(value).resolve()
            if self.storage_root in path.parents and path.is_file():
                path.unlink(missing_ok=True)

    def compare_versions(self, actor: dict[str, Any], left_id: str, right_id: str) -> dict[str, Any]:
        for version_id in (left_id, right_id):
            row = self.db.fetch_one(
                """SELECT v.version_id FROM teaching_archive_versions v JOIN courses c USING(course_id)
                   WHERE v.version_id=? AND c.owner_id=?""", (version_id, actor["user_id"]),
            )
            if not row:
                raise PermissionDenied("无权比较所选教学版本")
        rows = self.db.fetch_all(
            """SELECT version_id,record_type,title,structured_json,content_markdown
               FROM teaching_archive_items WHERE version_id IN (?,?) AND lifecycle!='withdrawn'""",
            (left_id, right_id),
        )
        sides: dict[str, dict[str, dict[str, Any]]] = {left_id: {}, right_id: {}}
        for row in rows:
            key = f"{row['record_type']}::{row['title']}"
            sides[row["version_id"]][key] = row
        differences = []
        for key in sorted(set(sides[left_id]) | set(sides[right_id])):
            left, right = sides[left_id].get(key), sides[right_id].get(key)
            state = "same"
            if left is None:
                state = "added"
            elif right is None:
                state = "removed"
            elif (left["structured_json"], left["content_markdown"]) != (right["structured_json"], right["content_markdown"]):
                state = "changed"
            differences.append({"key": key, "state": state,
                                "left": self._decode_item(left) if left else None,
                                "right": self._decode_item(right) if right else None})
        return {"left_version_id": left_id, "right_version_id": right_id, "differences": differences}

    def _require_document(self, actor: dict[str, Any], archive_document_id: str) -> tuple[dict[str, Any], Path]:
        row = self.db.fetch_one(
            """SELECT d.*,v.course_id FROM teaching_archive_documents d
               JOIN teaching_archive_versions v USING(version_id) JOIN courses c USING(course_id)
               WHERE d.archive_document_id=? AND c.owner_id=?""", (archive_document_id, actor["user_id"]),
        )
        if not row:
            raise PermissionDenied("无权访问该教学档案文件")
        source = Path(row.get("preview_path") or row["stored_path"]).resolve()
        campus_storage = self.campus.storage_dir.resolve()
        if (self.storage_root not in source.parents and campus_storage not in source.parents) or not source.is_file():
            raise NotFound("教学档案文件不存在")
        return row, source

    def preview_descriptor(self, actor: dict[str, Any], archive_document_id: str) -> dict[str, Any]:
        row, source = self._require_document(actor, archive_document_id)
        return {"archive_document_id": archive_document_id, "preview_kind": self._preview_kind(source),
                "conversion_status": row["conversion_status"],
                "preview_error": "" if row["conversion_status"] == "ready" else "旧版 Office 转换尚不可用"}

    @staticmethod
    def _xlsx_preview_html(source: Path) -> str:
        workbook = load_workbook(source, read_only=True, data_only=True)
        sections = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [html.escape(_clean(value, 500)) for value in row]
                if any(values):
                    rows.append("<tr>" + "".join(f"<td>{value}</td>" for value in values) + "</tr>")
                if len(rows) >= 500:
                    break
            sections.append(f"<h2>{html.escape(sheet.title)}</h2><div class='scroll'><table>{''.join(rows)}</table></div>")
        workbook.close()
        return """<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'><style>
        body{margin:0;padding:24px;background:#f3f6f8;color:#243447;font:14px/1.5 system-ui,'Microsoft YaHei'}
        h2{margin:20px 0 10px}.scroll{overflow:auto;background:white;border:1px solid #dde5ea;border-radius:10px}
        table{border-collapse:collapse;min-width:100%}td{padding:7px 10px;border:1px solid #e3e9ed;white-space:pre-wrap;min-width:90px}
        tr:first-child td{position:sticky;top:0;background:#e8f4f3;font-weight:700}</style></head><body>""" + "".join(sections) + "</body></html>"

    def preview_content(self, actor: dict[str, Any], archive_document_id: str) -> tuple[str, Path | str]:
        _row, source = self._require_document(actor, archive_document_id)
        suffix = source.suffix.lower()
        if suffix == ".docx":
            if not self.ingestion:
                raise NotFound("Word 预览服务不可用")
            return "text/html", self.ingestion._docx_preview_html(source)
        if suffix == ".xlsx":
            return "text/html", self._xlsx_preview_html(source)
        if suffix in {".txt", ".md"}:
            return "text/plain", source.read_text(encoding="utf-8-sig")
        if suffix == ".pdf":
            return "application/pdf", source
        if suffix == ".pptx":
            return "application/vnd.openxmlformats-officedocument.presentationml.presentation", source
        raise NotFound("该教学档案暂无可用预览")
