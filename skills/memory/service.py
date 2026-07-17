from __future__ import annotations

import io
import json
import re
from pathlib import Path
from typing import Any

from campus_service import CampusService, NotFound, PermissionDenied, ValidationError


def _loads(value: str | None) -> Any:
    return json.loads(value or "[]")


def _json_from_model(raw: str) -> Any:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.I | re.S)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise ValidationError(f"大模型返回格式不符合要求：{raw[:240]}") from exc


def _complete_json_objects(raw: str) -> list[dict]:
    """Recover complete JSON objects when a model response is cut off mid-array."""
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.I)
    objects: list[dict] = []
    start = None
    depth = 0
    in_string = False
    escaped = False
    for index, char in enumerate(cleaned):
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            if depth == 0:
                start = index
            depth += 1
        elif char == "}" and depth:
            depth -= 1
            if depth == 0 and start is not None:
                try:
                    value = json.loads(cleaned[start:index + 1])
                except json.JSONDecodeError:
                    pass
                else:
                    if isinstance(value, dict):
                        objects.append(value)
                start = None
    return objects


def _is_judgment_type(value: Any) -> bool:
    normalized = re.sub(r"[\s_\-/]+", "", str(value or "")).lower()
    return any(token in normalized for token in ("判断", "是非", "对错", "truefalse", "boolean"))


def _normalize_judgment_answer(value: Any) -> str:
    if isinstance(value, bool):
        return "正确" if value else "错误"
    if isinstance(value, (int, float)) and value in (0, 1):
        return "正确" if value == 1 else "错误"
    if isinstance(value, list):
        value = value[0] if value else ""
    normalized = str(value or "").strip()
    normalized = re.sub(r"^(?:正确)?答案\s*[:：]\s*", "", normalized, flags=re.I)
    normalized = normalized.strip(" \t\r\n()（）[]【】<>《》.。,:：;；")
    compact = re.sub(r"\s+", "", normalized).lower()
    true_values = {
        "对", "正确", "是", "真", "√", "✓", "✔", "t", "true", "yes", "y", "1", "right", "correct",
    }
    false_values = {
        "错", "错误", "否", "假", "×", "✕", "✖", "✗", "x", "f", "false", "no", "n", "0",
        "wrong", "incorrect",
    }
    if compact in true_values:
        return "正确"
    if compact in false_values:
        return "错误"
    return normalized


def _normalize_question_item(item: dict, default_type: str = "简答题") -> dict:
    item_type = str(item.get("type", default_type) or default_type).strip()
    question = str(item.get("question", "") or "").strip()
    options = [str(value).strip() for value in (item.get("options") or []) if str(value).strip()]
    answer = item.get("answer", "")
    if _is_judgment_type(item_type):
        if not answer:
            trailing = re.search(
                r"[（(]\s*(对|错|正确|错误|是|否|真|假|√|✓|✔|×|✕|✖|✗|T|F|True|False)\s*[）)]\s*$",
                question,
                flags=re.I,
            )
            if trailing:
                answer = trailing.group(1)
                question = question[:trailing.start()].rstrip()
        answer = _normalize_judgment_answer(answer)
        options = ["正确", "错误"]
        item_type = "判断题"
    elif isinstance(answer, list):
        answer = [str(value).strip() for value in answer if str(value).strip()]
    else:
        answer = str(answer or "").strip()
    return {
        "type":item_type,
        "question":question,
        "options":options,
        "answer":answer,
        "explanation":str(item.get("explanation", "") or "").strip(),
        "knowledge_point":str(item.get("knowledge_point", "课程知识点") or "课程知识点").strip(),
    }


class MemoryLearningSkill:
    """Student memory workflow: organize, chunk, train, evaluate and practice."""

    def __init__(self, campus: CampusService):
        self.campus = campus
        self.db = campus.db

    def _student_course(self, course_id: str, user_id: str) -> dict:
        course = self.campus.require_access(course_id, user_id, "student")
        return course

    def extract_image(self, course_id: str, user_id: str, file_name: str,
                      mime_type: str, image_bytes: bytes) -> dict:
        course = self._student_course(course_id, user_id)
        if course["course_type"] != "personal_course" or course["owner_id"] != user_id:
            raise PermissionDenied("图片资料只能保存到自己的个人课程")
        if mime_type not in {"image/png", "image/jpeg", "image/webp"}:
            raise ValidationError("图片仅支持 PNG、JPG 或 WEBP")
        if not image_bytes or len(image_bytes) > 10 * 1024 * 1024:
            raise ValidationError("图片不能为空且不能超过 10MB")
        provider = self.campus.provider_factory()
        extract_image_text = getattr(provider, "extract_image_text", None)
        if not callable(extract_image_text):
            raise ValidationError("当前智能接口不支持图片文字提取，请选择支持视觉输入的模型")
        text = extract_image_text(image_bytes, mime_type)
        if not text.strip():
            raise ValidationError("图片中没有提取到有效文字")
        safe_stem = Path(file_name).stem[:60] or "图片文字"
        document = self.campus.upload_document(
            course_id, user_id, "student", f"{safe_stem}_OCR.txt", "text/plain", text.encode("utf-8")
        )
        return {**document, "extracted_text": text}

    def build_blocks(self, course_id: str, user_id: str, document_id: str | None = None) -> list[dict]:
        self._student_course(course_id, user_id)
        params: tuple = (course_id,)
        sql = """SELECT c.content,c.section,c.document_id,d.original_name
                 FROM document_chunks c JOIN course_documents d USING(document_id)
                 WHERE c.course_id=?"""
        if document_id:
            sql += " AND c.document_id=?"
            params += (document_id,)
        rows = self.db.fetch_all(sql + " ORDER BY c.chunk_id", params)
        if not rows:
            raise ValidationError("当前资料没有可用于分块的后端文字")
        provider = self.campus.provider_factory()
        created: list[dict] = []
        for batch_start in range(0, len(rows), 8):
            batch = rows[batch_start:batch_start + 8]
            source_text = "\n\n".join(f"[{x['section']}]\n{x['content']}" for x in batch)[:14000]
            raw = provider.generate(
                "你是学习材料语义分块专家。只输出合法 JSON 数组，不要 Markdown。"
                "每项必须包含 title、keywords、content；title 是便于记忆的线索标题，"
                "keywords 是 3-8 个核心词，content 必须忠实保留原文信息，不得编造。",
                "请按逻辑段落和关键词密度把以下材料划分为知识块。每块适合独立背诵，"
                "标题尽量采用《主题的N个要点》这类记忆线索。\n\n" + source_text,
            )
            blocks = _json_from_model(raw)
            if not isinstance(blocks, list) or not blocks:
                raise ValidationError("大模型没有生成有效知识块")
            with self.db.connect() as conn:
                current = conn.execute("SELECT COALESCE(MAX(block_order),0) FROM knowledge_blocks WHERE course_id=?",
                                       (course_id,)).fetchone()[0]
                for offset, block in enumerate(blocks, 1):
                    if not isinstance(block, dict) or not str(block.get("content", "")).strip():
                        continue
                    title = str(block.get("title", "知识块")).strip()[:120]
                    keywords = [str(x).strip() for x in block.get("keywords", []) if str(x).strip()][:8]
                    cur = conn.execute("""INSERT INTO knowledge_blocks(course_id,document_id,owner_id,block_order,title,keywords_json,content)
                                        VALUES(?,?,?,?,?,?,?)""",
                                       (course_id, document_id or batch[0]["document_id"], user_id, current + offset,
                                        title, json.dumps(keywords, ensure_ascii=False), str(block["content"]).strip()))
                    created.append({"block_id": int(cur.lastrowid), "title": title, "keywords": keywords,
                                    "content": str(block["content"]).strip()})
        if not created:
            raise ValidationError("未能保存有效知识块")
        return created

    def list_blocks(self, course_id: str, user_id: str) -> list[dict]:
        self._student_course(course_id, user_id)
        rows = self.db.fetch_all("SELECT * FROM knowledge_blocks WHERE course_id=? ORDER BY block_order,block_id",
                                 (course_id,))
        for row in rows:
            row["keywords"] = _loads(row.pop("keywords_json"))
        return rows

    def _owned_block(self, block_id: int, user_id: str) -> dict:
        block = self.db.fetch_one("SELECT * FROM knowledge_blocks WHERE block_id=?", (block_id,))
        if not block:
            raise NotFound("知识块不存在")
        self._student_course(block["course_id"], user_id)
        if block["owner_id"] != user_id:
            raise PermissionDenied("只能调整自己创建的知识块")
        return block

    def update_block(self, block_id: int, user_id: str, title: str, keywords: list[str],
                     content: str, favorite: bool | None = None) -> dict:
        block = self._owned_block(block_id, user_id)
        if not title.strip() or not content.strip():
            raise ValidationError("标题和内容不能为空")
        favorite_value = block["is_favorite"] if favorite is None else int(favorite)
        self.db.execute("""UPDATE knowledge_blocks SET title=?,keywords_json=?,content=?,is_favorite=?,updated_at=CURRENT_TIMESTAMP
                         WHERE block_id=?""",
                        (title.strip()[:120], json.dumps([x.strip() for x in keywords if x.strip()][:12], ensure_ascii=False),
                         content.strip(), favorite_value, block_id))
        return next(x for x in self.list_blocks(block["course_id"], user_id) if x["block_id"] == block_id)

    def split_block(self, block_id: int, user_id: str, position: int) -> list[dict]:
        block = self._owned_block(block_id, user_id)
        content = block["content"]
        if position < 20 or position > len(content) - 20:
            raise ValidationError("拆分位置需保证前后知识块各至少 20 个字符")
        left, right = content[:position].strip(), content[position:].strip()
        with self.db.connect() as conn:
            conn.execute("UPDATE knowledge_blocks SET content=?,title=?,updated_at=CURRENT_TIMESTAMP WHERE block_id=?",
                         (left, f"{block['title']}（上）", block_id))
            conn.execute("UPDATE knowledge_blocks SET block_order=block_order+1 WHERE course_id=? AND block_order>?",
                         (block["course_id"], block["block_order"]))
            conn.execute("""INSERT INTO knowledge_blocks(course_id,document_id,owner_id,block_order,title,keywords_json,content)
                          VALUES(?,?,?,?,?,?,?)""",
                         (block["course_id"], block["document_id"], user_id, block["block_order"]+1,
                          f"{block['title']}（下）", block["keywords_json"], right))
        return self.list_blocks(block["course_id"], user_id)

    def merge_next(self, block_id: int, user_id: str) -> list[dict]:
        block = self._owned_block(block_id, user_id)
        next_block = self.db.fetch_one("""SELECT * FROM knowledge_blocks WHERE course_id=? AND block_order>?
                                        ORDER BY block_order,block_id LIMIT 1""",
                                       (block["course_id"], block["block_order"]))
        if not next_block or next_block["owner_id"] != user_id:
            raise ValidationError("没有可合并的下一个知识块")
        keywords = list(dict.fromkeys(_loads(block["keywords_json"]) + _loads(next_block["keywords_json"])))[:12]
        with self.db.connect() as conn:
            conn.execute("UPDATE knowledge_blocks SET content=?,keywords_json=?,updated_at=CURRENT_TIMESTAMP WHERE block_id=?",
                         (block["content"].rstrip()+"\n\n"+next_block["content"].lstrip(),
                          json.dumps(keywords, ensure_ascii=False), block_id))
            conn.execute("DELETE FROM knowledge_blocks WHERE block_id=?", (next_block["block_id"],))
        return self.list_blocks(block["course_id"], user_id)

    def cloze(self, block_id: int, user_id: str, extra_keywords: list[str] | None = None) -> dict:
        block = self._owned_block(block_id, user_id)
        keywords = list(dict.fromkeys(_loads(block["keywords_json"]) + (extra_keywords or [])))
        parts = self._cloze_parts(block["content"], keywords)
        if not any(x["type"] == "blank" for x in parts):
            raise ValidationError("关键词没有在知识块原文中出现，请先调整关键词")
        public_parts = [{k:v for k,v in item.items() if k != "answer"} for item in parts]
        return {"title": block["title"], "segments": public_parts,
                "blank_count": sum(x["type"] == "blank" for x in parts), "extra_keywords": extra_keywords or []}

    @staticmethod
    def _cloze_parts(content: str, keywords: list[str]) -> list[dict]:
        words = sorted({x.strip() for x in keywords if x.strip()}, key=len, reverse=True)
        if not words:
            return [{"type":"text", "value":content}]
        pattern = re.compile("|".join(re.escape(x) for x in words), re.I)
        parts, cursor, blank_index = [], 0, 0
        for match in pattern.finditer(content):
            if match.start() > cursor:
                parts.append({"type":"text", "value":content[cursor:match.start()]})
            blank_index += 1
            parts.append({"type":"blank", "index":blank_index, "length":len(match.group()), "answer":match.group()})
            cursor = match.end()
        if cursor < len(content):
            parts.append({"type":"text", "value":content[cursor:]})
        return parts

    def submit_cloze(self, block_id: int, user_id: str, extra_keywords: list[str],
                     responses: list[str]) -> dict:
        block = self._owned_block(block_id, user_id)
        keywords = list(dict.fromkeys(_loads(block["keywords_json"]) + (extra_keywords or [])))
        blanks = [x for x in self._cloze_parts(block["content"], keywords) if x["type"] == "blank"]
        if not blanks:
            raise ValidationError("当前知识块没有可检测的挖空")
        corrections, missing, errors = [], [], []
        correct_count = 0
        for index, blank in enumerate(blanks):
            response = str(responses[index] if index < len(responses) else "").strip()
            expected = blank["answer"]
            correct = re.sub(r"\s+", "", response).lower() == re.sub(r"\s+", "", expected).lower()
            correct_count += int(correct)
            if not response:
                missing.append(expected)
            elif not correct:
                errors.append(f"第 {index+1} 空：填写“{response}”，正确答案“{expected}”")
            corrections.append({"index":index+1,"response":response,"correct_answer":expected,"correct":correct})
        score = round(100 * correct_count / len(blanks), 1)
        feedback = "全部填写正确。" if score == 100 else f"共 {len(blanks)} 空，答对 {correct_count} 空；错误已加入背诵本。"
        attempt_id = self.db.execute("""INSERT INTO memory_attempts(course_id,block_id,user_id,mode,score,missing_points_json,error_points_json,feedback)
                                    VALUES(?,?,?,?,?,?,?,?)""",
                                   (block["course_id"],block_id,user_id,"cloze",score,
                                    json.dumps(missing,ensure_ascii=False),json.dumps(errors,ensure_ascii=False),feedback))
        return {"attempt_id":attempt_id,"score":score,"correct_count":correct_count,"total":len(blanks),
                "corrections":corrections,"missing_points":missing,"error_points":errors,"feedback":feedback}

    def evaluate_recitation(self, block_id: int, user_id: str, recited_text: str) -> dict:
        block = self._owned_block(block_id, user_id)
        if len(recited_text.strip()) < 5:
            raise ValidationError("请输入或粘贴背诵内容后再检测")
        provider = self.campus.provider_factory()
        raw = provider.generate(
            "你是严格但友善的背诵监督员。只输出 JSON 对象，字段为 score(0-100)、"
            "missing_points(数组)、error_points(数组)、feedback(字符串)。不得补充原文外知识。",
            f"原始知识块：\n{block['content']}\n\n学生背诵：\n{recited_text.strip()}",
        )
        result = _json_from_model(raw)
        if not isinstance(result, dict):
            raise ValidationError("大模型未返回有效检测结果")
        score = max(0.0, min(100.0, float(result.get("score", 0))))
        missing = [str(x) for x in result.get("missing_points", [])]
        errors = [str(x) for x in result.get("error_points", [])]
        feedback = str(result.get("feedback", ""))
        attempt_id = self.db.execute("""INSERT INTO memory_attempts(course_id,block_id,user_id,mode,score,missing_points_json,error_points_json,feedback)
                                    VALUES(?,?,?,?,?,?,?,?)""",
                                   (block["course_id"], block_id, user_id, "recitation", score,
                                    json.dumps(missing, ensure_ascii=False), json.dumps(errors, ensure_ascii=False), feedback))
        return {"attempt_id": attempt_id, "score": score, "missing_points": missing,
                "error_points": errors, "feedback": feedback}

    def memory_summary(self, course_id: str, user_id: str) -> dict:
        self._student_course(course_id, user_id)
        rows = self.db.fetch_all("""SELECT a.*,b.title FROM memory_attempts a LEFT JOIN knowledge_blocks b USING(block_id)
                                  WHERE a.course_id=? AND a.user_id=? ORDER BY a.attempt_id DESC""",
                                 (course_id, user_id))
        for row in rows:
            row["missing_points"] = _loads(row.pop("missing_points_json"))
            row["error_points"] = _loads(row.pop("error_points_json"))
        avg = round(sum(x["score"] for x in rows) / len(rows), 1) if rows else 0
        weak = []
        for row in rows:
            weak.extend(row["missing_points"] + row["error_points"])
        counts: dict[str, int] = {}
        for point in weak:
            counts[point] = counts.get(point, 0) + 1
        return {"attempts": rows, "average_score": avg,
                "weak_points": [{"point": k, "count": v} for k,v in sorted(counts.items(), key=lambda x:-x[1])[:10]]}

    def generate_questions(self, course_id: str, user_id: str, count: int = 6) -> list[dict]:
        blocks = self.list_blocks(course_id, user_id)
        if not blocks:
            raise ValidationError("请先生成知识块")
        source = "\n\n".join(f"【{x['title']}】\n{x['content']}" for x in blocks)[:16000]
        raw = self.campus.provider_factory().generate(
            "你是课程出题专家。只输出 JSON 数组。题型必须覆盖单选题、多选题、判断题、简答题。"
            "每题字段：type、question、options（简答题为空）、answer（多选题为答案数组，其余为字符串）、"
            "explanation、knowledge_point。单选/判断题只能有一个正确答案，多选题至少两个正确答案。题目只能依据材料。",
            f"请依据以下学习材料生成 {max(3,min(count,12))} 道难度递进的练习题：\n\n{source}",
        )
        questions = _json_from_model(raw)
        if not isinstance(questions, list) or not questions:
            raise ValidationError("大模型没有生成有效题目")
        cleaned = []
        for item in questions[:12]:
            if isinstance(item, dict) and item.get("question") and "answer" in item:
                cleaned.append(_normalize_question_item(item))
        if not cleaned:
            raise ValidationError("大模型题目格式无效")
        self.db.execute("INSERT INTO generated_practice(course_id,user_id,questions_json) VALUES(?,?,?)",
                        (course_id,user_id,json.dumps(cleaned,ensure_ascii=False)))
        return cleaned

    @staticmethod
    def _question_bank_text(file_name: str, data: bytes) -> str:
        suffix = Path(file_name).suffix.lower()
        if suffix == ".txt":
            for encoding in ("utf-8-sig", "gb18030"):
                try:
                    return data.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValidationError("TXT 题库编码无法识别，请保存为 UTF-8 后重试")
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join(
                f"【第 {index} 页】\n{page.extract_text() or ''}"
                for index, page in enumerate(reader.pages, 1)
            )
        if suffix == ".docx":
            from docx import Document
            document = Document(io.BytesIO(data))
            lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table_index, table in enumerate(document.tables, 1):
                lines.append(f"【表格 {table_index}】")
                lines.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            return "\n".join(lines)
        if suffix == ".xlsx":
            from openpyxl import load_workbook
            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"【工作表：{sheet.title}】")
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if value is None else str(value).strip() for value in row]
                    if any(values):
                        lines.append("\t".join(values))
            workbook.close()
            return "\n".join(lines)
        if suffix == ".xls":
            raise ValidationError("暂不支持旧版 .xls，请在 Excel 中另存为 .xlsx 后导入")
        raise ValidationError("题库仅支持 PDF、DOCX、TXT 和 Excel（XLSX）")

    def import_question_bank(self, course_id: str, user_id: str, file_name: str,
                             mime_type: str, data: bytes) -> list[dict]:
        self._student_course(course_id, user_id)
        suffix = Path(file_name).suffix.lower()
        allowed_mimes = {
            ".pdf":{"application/pdf", "application/octet-stream"},
            ".docx":{"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/octet-stream"},
            ".txt":{"text/plain", "application/octet-stream"},
            ".xlsx":{"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/octet-stream"},
            ".xls":{"application/vnd.ms-excel", "application/octet-stream"},
        }
        if suffix not in allowed_mimes or mime_type not in allowed_mimes[suffix]:
            raise ValidationError("题库文件扩展名与文件类型不匹配")
        if not data or len(data) > 20 * 1024 * 1024:
            raise ValidationError("题库文件不能为空且不能超过 20MB")
        source_text = self._question_bank_text(file_name, data).strip()
        if len(source_text) < 3:
            raise ValidationError("题库中没有提取到可识别的文字；扫描版 PDF 请先进行 OCR")
        provider = self.campus.provider_factory()
        system_prompt = (
            "你是题库结构化整理智能体。只输出合法 JSON 数组，不要输出 Markdown。"
            "逐题提取 type、question、options、answer、explanation、knowledge_point。"
            "type 统一为单选题、多选题、判断题或简答题；没有选项时 options 为空数组。"
            "题库明确给出答案时必须原样提取；未给答案时 answer 必须是空字符串，严禁编造答案。"
            "多选题已有答案时 answer 使用字符串数组，其余题型使用字符串。"
            "每次最多提取 20 题，确保 JSON 完整闭合；不要输出未完成的题目。"
        )
        source_chunks = self._question_bank_chunks(source_text[:100000])
        imported: list[dict] = []
        for chunk_index, chunk in enumerate(source_chunks, 1):
            imported.extend(self._parse_question_bank_chunk(
                provider, system_prompt, Path(file_name).name, chunk, chunk_index
            ))
            if len(imported) >= 100:
                break
        if not imported:
            raise ValidationError("AI 未能从题库中识别出有效题目")
        questions = []
        seen_questions: set[str] = set()
        for item in imported:
            if not isinstance(item, dict) or not str(item.get("question", "")).strip():
                continue
            question_text = str(item["question"]).strip()
            dedupe_key = re.sub(r"\s+", "", question_text).lower()
            if dedupe_key in seen_questions:
                continue
            seen_questions.add(dedupe_key)
            normalized_item = _normalize_question_item({**item, "question":question_text}, "简答题")
            normalized_item["answer_source"] = "imported" if normalized_item["answer"] else "ai_judge"
            normalized_item["source_file"] = Path(file_name).name
            questions.append(normalized_item)
            if len(questions) >= 100:
                break
        if not questions:
            raise ValidationError("题库中没有可保存的有效题目")
        self.db.execute("INSERT INTO generated_practice(course_id,user_id,questions_json) VALUES(?,?,?)",
                        (course_id, user_id, json.dumps(questions, ensure_ascii=False)))
        return questions

    @staticmethod
    def _question_bank_chunks(text: str, target_size: int = 4500, overlap: int = 320) -> list[str]:
        if len(text) <= target_size:
            return [text]
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + target_size, len(text))
            if end < len(text):
                boundary = max(text.rfind("\n\n", start + target_size // 2, end),
                               text.rfind("\n", start + target_size // 2, end))
                if boundary > start:
                    end = boundary
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start = max(end - overlap, start + 1)
        return chunks

    def _parse_question_bank_chunk(self, provider, system_prompt: str, file_name: str,
                                   chunk: str, chunk_index: int, depth: int = 0) -> list[dict]:
        raw = provider.generate(
            system_prompt,
            f"文件名：{file_name}\n当前是第 {chunk_index} 个片段。"
            f"\n\n请只整理这个片段中的完整题目：\n{chunk}",
        )
        try:
            value = _json_from_model(raw)
            if isinstance(value, list):
                return [item for item in value if isinstance(item, dict)]
        except ValidationError:
            if depth < 3 and len(chunk) > 900:
                midpoint = len(chunk) // 2
                split_at = chunk.rfind("\n", 0, midpoint)
                if split_at < len(chunk) // 4:
                    split_at = midpoint
                left = chunk[:split_at].strip()
                right = chunk[split_at:].strip()
                return (
                    self._parse_question_bank_chunk(
                        provider, system_prompt, file_name, left, chunk_index, depth + 1
                    )
                    + self._parse_question_bank_chunk(
                        provider, system_prompt, file_name, right, chunk_index, depth + 1
                    )
                )
            recovered = _complete_json_objects(raw)
            if recovered:
                return recovered
            raise
        raise ValidationError("AI 返回的题库片段不是 JSON 数组")

    def grade_questions(self, course_id: str, user_id: str, questions: list[dict],
                        responses: list[Any]) -> dict:
        self._student_course(course_id, user_id)
        if not questions or len(responses) != len(questions):
            raise ValidationError("题目或作答数据不完整")
        normalized_questions = []
        normalized_responses = []
        for question, response in zip(questions, responses):
            normalized = _normalize_question_item(question)
            normalized["answer_source"] = question.get(
                "answer_source", "imported" if normalized["answer"] else "ai_judge"
            )
            for key in ("source_file",):
                if key in question:
                    normalized[key] = question[key]
            normalized_questions.append(normalized)
            normalized_responses.append(
                _normalize_judgment_answer(response) if _is_judgment_type(normalized["type"]) else response
            )
        grading_input = [{"index":i+1,"type":q.get("type"),"question":q.get("question"),
                          "standard_answer":q.get("answer"),"explanation":q.get("explanation", ""),
                          "answer_source":q.get("answer_source", "imported" if q.get("answer") else "ai_judge"),
                          "options":q.get("options", []),
                          "knowledge_point":q.get("knowledge_point", ""),"student_answer":normalized_responses[i]}
                         for i,q in enumerate(normalized_questions)]
        raw = self.campus.provider_factory().generate(
            "你是严谨的课程练习批改智能体。只输出 JSON 对象，字段：score（0-100）、results（数组，"
            "每项含 index、correct、correct_answer、feedback）、weak_points（数组）、summary。"
            "standard_answer 非空时按标准答案批改，简答题允许语义等价。"
            "standard_answer 为空或 answer_source=ai_judge 时，必须先独立求解题目，再判断学生答案，"
            "并在 correct_answer 中给出你判定的正确答案；不能因为学生填写了某答案就默认正确。",
            "请批改以下作答：\n" + json.dumps(grading_input, ensure_ascii=False),
        )
        result = _json_from_model(raw)
        if not isinstance(result, dict) or not isinstance(result.get("results"), list):
            raise ValidationError("大模型未返回有效批改结果")
        for result_item in result["results"]:
            try:
                question_index = int(result_item.get("index", 0)) - 1
            except (TypeError, ValueError):
                continue
            if 0 <= question_index < len(normalized_questions):
                question = normalized_questions[question_index]
                if _is_judgment_type(question["type"]):
                    result_item["correct_answer"] = _normalize_judgment_answer(
                        result_item.get("correct_answer", question.get("answer", ""))
                    )
        score = max(0.0, min(100.0, float(result.get("score", 0))))
        result["score"] = score
        attempt_id = self.db.execute("""INSERT INTO ai_practice_attempts(course_id,user_id,questions_json,responses_json,result_json,score)
                                    VALUES(?,?,?,?,?,?)""",
                                   (course_id,user_id,json.dumps(normalized_questions,ensure_ascii=False),
                                    json.dumps(normalized_responses,ensure_ascii=False),
                                    json.dumps(result,ensure_ascii=False),score))
        result["attempt_id"] = attempt_id
        return result

    def student_dashboard(self, user_id: str, course_id: str | None = None) -> dict:
        courses = self.campus.list_courses(user_id, "student")
        if course_id:
            self._student_course(course_id, user_id)
            course_ids = [course_id]
        else:
            course_ids = [x["course_id"] for x in courses]
        if not course_ids:
            return {"course_count":0,"document_count":0,"block_count":0,"memory_attempts":[],
                    "practice_attempts":[],"memory_average":0,"practice_average":0,"recitation_book":[],
                    "wrong_question_book":[],"weak_points":[]}
        placeholders = ",".join("?" for _ in course_ids)
        documents = self.db.fetch_one(f"SELECT COUNT(*) count FROM course_documents WHERE course_id IN ({placeholders})", tuple(course_ids))
        blocks = self.db.fetch_one(f"SELECT COUNT(*) count FROM knowledge_blocks WHERE course_id IN ({placeholders}) AND owner_id=?",
                                   tuple(course_ids)+(user_id,))
        memory = self.db.fetch_all(f"""SELECT a.*,b.title FROM memory_attempts a LEFT JOIN knowledge_blocks b USING(block_id)
                                    WHERE a.course_id IN ({placeholders}) AND a.user_id=? ORDER BY a.attempt_id DESC""",
                                   tuple(course_ids)+(user_id,))
        practices = self.db.fetch_all(f"""SELECT * FROM ai_practice_attempts WHERE course_id IN ({placeholders}) AND user_id=?
                                       ORDER BY attempt_id DESC""", tuple(course_ids)+(user_id,))
        recitation_book, point_counts = [], {}
        for row in memory:
            row["missing_points"] = _loads(row.pop("missing_points_json"))
            row["error_points"] = _loads(row.pop("error_points_json"))
            if row["missing_points"] or row["error_points"] or row["score"] < 100:
                recitation_book.append(row)
            for point in row["missing_points"] + row["error_points"]:
                point_counts[point] = point_counts.get(point, 0) + 1
        for row in practices:
            row["questions"] = _loads(row.pop("questions_json")); row["responses"] = _loads(row.pop("responses_json"))
            row["result"] = json.loads(row.pop("result_json") or "{}")
            for point in row["result"].get("weak_points", []):
                point_counts[str(point)] = point_counts.get(str(point), 0) + 1
        wrong_question_book = []
        for attempt in practices:
            for result_item in attempt["result"].get("results", []):
                if result_item.get("correct"):
                    continue
                index = int(result_item.get("index", 0)) - 1
                if index < 0 or index >= len(attempt["questions"]):
                    continue
                question = attempt["questions"][index]
                response = attempt["responses"][index] if index < len(attempt["responses"]) else ""
                wrong_question_book.append({
                    "question":question.get("question", ""), "type":question.get("type", ""),
                    "student_answer":response, "correct_answer":result_item.get("correct_answer", question.get("answer", "")),
                    "feedback":result_item.get("feedback", ""),
                    "knowledge_point":question.get("knowledge_point", ""), "created_at":attempt["created_at"],
                })
        return {
            "course_count":len(courses), "document_count":documents["count"] if documents else 0,
            "block_count":blocks["count"] if blocks else 0, "memory_attempts":memory,"practice_attempts":practices,
            "memory_average":round(sum(x["score"] for x in memory)/len(memory),1) if memory else 0,
            "practice_average":round(sum(x["score"] for x in practices)/len(practices),1) if practices else 0,
            "recitation_book":recitation_book,
            "wrong_question_book":wrong_question_book,
            "weak_points":[{"point":k,"count":v} for k,v in sorted(point_counts.items(),key=lambda x:-x[1])[:15]],
        }

    @staticmethod
    def export_workbook(course_name: str, questions: list[dict]) -> bytes:
        from docx import Document
        document = Document(); document.add_heading(f"{course_name} 练习册", 0)
        for index, item in enumerate(questions, 1):
            document.add_paragraph(f"{index}. [{item['type']}] {item['question']}")
            for option in item.get("options", []):
                document.add_paragraph(option, style="List Bullet")
        document.add_page_break(); document.add_heading("参考答案与解析", 0)
        for index, item in enumerate(questions, 1):
            document.add_paragraph(f"{index}. {item['answer']}")
            if item.get("explanation"):
                document.add_paragraph(f"解析：{item['explanation']}")
        MemoryLearningSkill._apply_songti(document)
        output = io.BytesIO(); document.save(output); return output.getvalue()

    @staticmethod
    def _apply_songti(document) -> None:
        from docx.oxml.ns import qn
        from docx.shared import Pt

        def set_run_font(run) -> None:
            run.font.name = "宋体"
            run_properties = run._element.get_or_add_rPr()
            run_fonts = run_properties.get_or_add_rFonts()
            run_fonts.set(qn("w:ascii"), "SimSun")
            run_fonts.set(qn("w:hAnsi"), "SimSun")
            run_fonts.set(qn("w:eastAsia"), "宋体")

        for style in document.styles:
            if hasattr(style, "font"):
                style.font.name = "宋体"
                style.font.size = style.font.size or Pt(11)
                run_properties = style.element.get_or_add_rPr()
                run_fonts = run_properties.get_or_add_rFonts()
                run_fonts.set(qn("w:ascii"), "SimSun")
                run_fonts.set(qn("w:hAnsi"), "SimSun")
                run_fonts.set(qn("w:eastAsia"), "宋体")
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                set_run_font(run)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run)

    def export_recitation_book(self, course_id: str, user_id: str, course_name: str) -> bytes:
        from docx import Document
        dashboard = self.student_dashboard(user_id, course_id)
        document = Document(); document.add_heading(f"{course_name} · 个人背诵本", 0)
        if not dashboard["recitation_book"]:
            document.add_paragraph("暂无背诵错误记录。")
        for index, item in enumerate(dashboard["recitation_book"], 1):
            document.add_heading(f"{index}. {item.get('title') or '知识块'}", level=1)
            document.add_paragraph(f"训练方式：{item['mode']}　成绩：{item['score']:.1f}%　时间：{item['created_at']}")
            if item["missing_points"]:
                document.add_paragraph("缺失内容：" + "；".join(item["missing_points"]))
            if item["error_points"]:
                document.add_paragraph("错误内容：" + "；".join(item["error_points"]))
            document.add_paragraph("学习建议：" + item["feedback"])
        self._apply_songti(document)
        output = io.BytesIO(); document.save(output); return output.getvalue()

    def export_wrong_question_book(self, course_id: str, user_id: str, course_name: str) -> bytes:
        from docx import Document
        dashboard = self.student_dashboard(user_id, course_id)
        document = Document(); document.add_heading(f"{course_name} · 个人错题本", 0)
        if not dashboard["wrong_question_book"]:
            document.add_paragraph("暂无错题记录。")
        for index, item in enumerate(dashboard["wrong_question_book"], 1):
            document.add_heading(f"{index}. [{item['type']}] {item['question']}", level=1)
            student_answer = item["student_answer"]
            if isinstance(student_answer, list): student_answer = "、".join(map(str,student_answer))
            correct_answer = item["correct_answer"]
            if isinstance(correct_answer, list): correct_answer = "、".join(map(str,correct_answer))
            document.add_paragraph(f"我的答案：{student_answer or '未作答'}")
            document.add_paragraph(f"正确答案：{correct_answer}")
            document.add_paragraph(f"知识点：{item['knowledge_point'] or '未标注'}")
            document.add_paragraph(f"解析：{item['feedback']}")
        self._apply_songti(document)
        output = io.BytesIO(); document.save(output); return output.getvalue()
