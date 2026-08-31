import io
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from auth_service import AuthService
from campus_service import CampusService
from database import LearningDatabase
from teacher_service import TeacherService
from teaching_archive_service import TeachingArchiveService


def archive_scope(tmp_path: Path):
    db = LearningDatabase(tmp_path / "archive.db")
    campus = CampusService(db, tmp_path / "uploads", provider_factory=lambda: None)
    teacher = AuthService(db, tmp_path / "secret").create_user(
        "archive-teacher", "safe-password-123", "teacher"
    )
    teachers = TeacherService(db, campus)
    course = teachers.create_course(teacher, "数据库原理与应用")
    term = teachers.create_term(
        teacher, "2025-2026 第一学期", academic_year="2025-2026", teaching_period="第一学期"
    )
    class_row = teachers.create_class(
        teacher, course["course_id"], term["term_id"], "仁24信管1", "A班",
        "周三3-5节", "仁济", "2024", "信息管理与信息系统", "标准",
    )
    return db, teacher, course, term, class_row, TeachingArchiveService(db, campus)


def lesson_plan_bytes() -> bytes:
    document = Document()
    document.add_heading("温州医科大学教案", level=1)
    table = document.add_table(rows=4, cols=2)
    for row, values in zip(table.rows, (
        ("课程名称", "数据库原理与应用"),
        ("授课对象", "仁济24级信管"),
        ("授课教师", "刘老师"),
        ("教师单位", "温州医科大学"),
    )):
        for cell, value in zip(row.cells, values):
            cell.text = value
    design = document.add_table(rows=6, cols=2)
    for row, values in zip(design.rows, (
        ("授课章节题目", "第1章 数据库系统概述"),
        ("授课时长", "3课时"),
        ("教学方式", "线下面授"),
        ("学情分析", "具备计算机基础"),
        ("教学目标", "理解数据库基本概念"),
        ("教学反思", "后续增加案例"),
    )):
        row.cells[0].text, row.cells[1].text = values
    steps = document.add_table(rows=2, cols=7)
    for index, value in enumerate(("教学步骤", "教学环节", "教学内容", "教师活动", "学生活动", "设计意图", "时间")):
        steps.rows[0].cells[index].text = value
    for index, value in enumerate(("课前", "预习", "基本概念", "提出问题", "讨论", "激活旧知", "10分钟")):
        steps.rows[1].cells[index].text = value
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def schedule_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "仁24信管"
    sheet.append(["温州医科大学教学进度表"])
    sheet.append(["周次", "日期", "星期", "节次", "课程名称", "授课班级", "理论/实验学时", "理论/实验授课内容", "教师", "授课性质"])
    sheet.append([1, "2025-09-10", "三", "3-5", "数据库原理与应用", "仁24信管1", 3, "数据库系统概述", "刘老师", "理论"])
    sheet.append([2, "2025-09-17", "三", "8-11", "数据库原理与应用", "仁24信管1", 4, "SQL管理器简单应用", "刘老师", "实验"])
    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def experiment_syllabus_bytes() -> bytes:
    document = Document()
    document.add_heading("实验教学大纲", level=1)
    table = document.add_table(rows=2, cols=5)
    for index, value in enumerate(("序号", "实验项目名称", "学时", "要求", "地点")):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(("1", "SQL server 2016管理器简单应用", "4", "必做", "机房")):
        table.rows[1].cells[index].text = value
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def archive_bytes() -> bytes:
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("readme.txt", "实验材料")
    return stream.getvalue()


def test_class_dimensions_and_lesson_plan_batch_publish(tmp_path: Path):
    _db, teacher, course, term, class_row, service = archive_scope(tmp_path)
    assert class_row["campus"] == "仁济"
    assert class_row["cohort_year"] == "2024"
    assert class_row["major"] == "信息管理与信息系统"
    batch = service.create_import_batch(
        teacher, course["course_id"], term_id=term["term_id"],
        defaults={"campus": "仁济", "cohort_year": "2024", "major": "信息管理与信息系统"},
    )
    uploaded = service.add_import_file(
        teacher, batch["batch_id"], "数据库原理与应用-第1次课-教案.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        io.BytesIO(lesson_plan_bytes()), relative_path="教案/数据库原理与应用-第1次课-教案.docx",
    )
    service.update_import_file(
        teacher, batch["batch_id"], uploaded["file_id"], {"class_ids": [class_row["class_id"]]},
    )
    committed = service.commit_import_batch(teacher, batch["batch_id"])
    assert committed["status"] == "completed"
    workbench = service.workbench(teacher, course["course_id"])
    lesson = next(item for item in workbench["items"] if item["record_type"] == "lesson_session")
    assert lesson["structured"]["session_number"] == 1
    assert lesson["structured"]["chapter"] == "第1章 数据库系统概述"
    assert lesson["structured"]["teaching_steps"][1][2] == "基本概念"
    assert any(item["record_type"] == "teaching_reflection" for item in workbench["items"])


def test_schedule_and_assessment_quality_gate(tmp_path: Path):
    _db, teacher, course, term, class_row, service = archive_scope(tmp_path)
    batch = service.create_import_batch(teacher, course["course_id"], term_id=term["term_id"])
    schedule = service.add_import_file(
        teacher, batch["batch_id"], "2025-2026教学进度.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        io.BytesIO(schedule_bytes()), relative_path="教学进度/仁济/信管/2025-2026教学进度.xlsx",
    )
    assessment = service.add_import_file(
        teacher, batch["batch_id"], "考核标准及内容.txt", "text/plain",
        io.BytesIO("平时成绩20%+阶段考核成绩30%+期末成绩40%".encode()),
    )
    for item in (schedule, assessment):
        service.update_import_file(
            teacher, batch["batch_id"], item["file_id"], {"class_ids": [class_row["class_id"]]},
        )
    committed = service.commit_import_batch(teacher, batch["batch_id"])
    assert committed["status"] == "completed_with_warnings"
    workbench = service.workbench(teacher, course["course_id"])
    schedule_items = [item for item in workbench["items"] if item["record_type"] == "teaching_schedule_entry"]
    assert len(schedule_items) == 2
    assessment_item = next(item for item in workbench["items"] if item["record_type"] == "assessment_scheme")
    assert assessment_item["lifecycle"] == "review_required"
    assert "assessment_weight_mismatch" in assessment_item["risk_codes"]
    components = [item for item in workbench["items"] if item["record_type"] == "assessment_component"]
    assert [item["structured"]["weight"] for item in components] == [20.0, 30.0, 40.0]
    assert all(item["lifecycle"] == "review_required" for item in components)


def test_executable_is_blocked_without_aborting_batch(tmp_path: Path):
    _db, teacher, course, term, _class_row, service = archive_scope(tmp_path)
    batch = service.create_import_batch(teacher, course["course_id"], term_id=term["term_id"])
    blocked = service.add_import_file(
        teacher, batch["batch_id"], "登录注册.exe", "application/octet-stream",
        io.BytesIO(b"MZ" + b"0" * 30), relative_path="实验/实验4/登录注册.exe",
    )
    assert blocked["status"] == "blocked"
    assert blocked["risk_codes"] == ["executable_content_blocked"]
    committed = service.commit_import_batch(teacher, batch["batch_id"])
    assert committed["error_count"] == 1


def test_experiment_attachment_is_linked_by_number_and_exposes_source_path(tmp_path: Path):
    _db, teacher, course, term, class_row, service = archive_scope(tmp_path)
    batch = service.create_import_batch(teacher, course["course_id"], term_id=term["term_id"])
    syllabus = service.add_import_file(
        teacher, batch["batch_id"], "实验教学大纲.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        io.BytesIO(experiment_syllabus_bytes()), relative_path="大纲/实验教学大纲.docx",
    )
    attachment = service.add_import_file(
        teacher, batch["batch_id"], "实验一源码.zip", "application/zip",
        io.BytesIO(archive_bytes()), relative_path="实验/实验一/实验一源码.zip",
    )
    service.update_import_file(teacher, batch["batch_id"], syllabus["file_id"], {
        "class_ids": [class_row["class_id"]], "record_type": "syllabus",
        "routing_target": "teaching_archive",
    })
    service.update_import_file(teacher, batch["batch_id"], attachment["file_id"], {
        "class_ids": [class_row["class_id"]], "routing_target": "attachment",
    })
    service.commit_import_batch(teacher, batch["batch_id"])
    workbench = service.workbench(teacher, course["course_id"])
    project = next(item for item in workbench["items"] if item["record_type"] == "experiment_project")
    linked = workbench["attachments"][0]
    assert linked["experiment_item_id"] == project["item_id"]
    assert linked["relative_path"] == "实验/实验一/实验一源码.zip"


def test_duplicate_defaults_to_skip_and_requires_explicit_new_version(tmp_path: Path):
    _db, teacher, course, term, class_row, service = archive_scope(tmp_path)
    content = lesson_plan_bytes()
    first = service.create_import_batch(teacher, course["course_id"], term_id=term["term_id"])
    uploaded = service.add_import_file(
        teacher, first["batch_id"], "第1次课.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        io.BytesIO(content), relative_path="教案/第1次课.docx",
    )
    service.update_import_file(teacher, first["batch_id"], uploaded["file_id"], {
        "class_ids": [class_row["class_id"]],
    })
    service.commit_import_batch(teacher, first["batch_id"])

    duplicate_batch = service.create_import_batch(
        teacher, course["course_id"], term_id=term["term_id"],
    )
    duplicate = service.add_import_file(
        teacher, duplicate_batch["batch_id"], "第1次课-副本.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        io.BytesIO(content), relative_path="教案/第1次课-副本.docx",
    )
    assert duplicate["duplicate_action"] == "skip"
    assert "duplicate_existing_document" in duplicate["risk_codes"]
    service.commit_import_batch(teacher, duplicate_batch["batch_id"])
    assert len(service.workbench(teacher, course["course_id"])["documents"]) == 1

    version_batch = service.create_import_batch(
        teacher, course["course_id"], term_id=term["term_id"],
    )
    version_file = service.add_import_file(
        teacher, version_batch["batch_id"], "第1次课-v2.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        io.BytesIO(content), relative_path="教案/第1次课-v2.docx",
    )
    service.update_import_file(teacher, version_batch["batch_id"], version_file["file_id"], {
        "class_ids": [class_row["class_id"]], "duplicate_action": "new_version",
    })
    service.commit_import_batch(teacher, version_batch["batch_id"])
    workbench = service.workbench(teacher, course["course_id"])
    assert len(workbench["documents"]) == 2
    assert any(version["version_number"] == 2 for version in workbench["versions"])
    deleted = service.delete_document(teacher, workbench["documents"][0]["archive_document_id"])
    assert deleted["deleted"]
    assert len(service.workbench(teacher, course["course_id"])["documents"]) == 1
